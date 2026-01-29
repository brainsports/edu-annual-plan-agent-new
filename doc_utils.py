import io
import os
import requests
import pandas as pd
import matplotlib
import logging

matplotlib.use("Agg")
import matplotlib.pyplot as plt
import matplotlib.font_manager as fm
import numpy as np

from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

SURVEY_QUESTION_COUNT = 10
EXPECTED_EFFECT_MAX_CHARS = 100  # 공백 포함 100자

logger = logging.getLogger(__name__)


def truncate_expected_effect(text: str, max_chars: int = EXPECTED_EFFECT_MAX_CHARS) -> str:
    """기대효과 100자 제한 (공백 포함). 축약 후에도 최소 10자 이상 유지."""
    if not text:
        return text
    
    text = str(text).strip()
    original_len = len(text)
    
    if len(text) <= max_chars:
        return text
    
    # 축약: 마지막 완성된 문장에서 자르기 시도
    truncated = text[:max_chars]
    
    # 문장 끝 기호 찾기
    last_period = max(truncated.rfind('.'), truncated.rfind('다.'), truncated.rfind('함.'), truncated.rfind('됨.'))
    if last_period > max_chars // 2:
        truncated = truncated[:last_period + 1]
    else:
        # 단어 경계에서 자르기
        last_space = truncated.rfind(' ')
        if last_space > max_chars // 2:
            truncated = truncated[:last_space]
    
    if len(truncated) < 10:
        truncated = text[:max_chars]
    
    logger.debug(f"[기대효과 100자] {original_len}→{len(truncated)} chars")
    return truncated

# --- 기본 서식 함수 ---
def ensure_korean_font():
    """
    NanumGothic 폰트를 다운로드하여 matplotlib용 FontProperties로 반환.
    (docx에 직접 적용하는 기능은 아니며, 차트 생성 시 한글깨짐 방지에 사용 가능)
    """
    font_filename = "NanumGothic.ttf"
    font_url = "https://github.com/google/fonts/raw/main/ofl/nanumgothic/NanumGothic-Regular.ttf"
    if not os.path.exists(font_filename):
        try:
            response = requests.get(font_url, timeout=10)
            response.raise_for_status()
            with open(font_filename, "wb") as f:
                f.write(response.content)
        except Exception:
            return None

    try:
        return fm.FontProperties(fname=font_filename)
    except Exception:
        return None


def set_standard_margins(document: Document):
    for section in document.sections:
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)


def set_cell_background(cell, color_hex: str):
    """테이블 셀 배경색 설정 (예: "D9D9D9" = 연한 회색)"""
    shading_elm = parse_xml(
        r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), color_hex)
    )
    cell._tc.get_or_add_tcPr().append(shading_elm)


def add_left_aligned_heading(document: Document, text: str, level: int = 1):
    heading = document.add_heading(text, level=level)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return heading


def add_markdown_text(cell_or_paragraph, text: str):
    """
    **bold** 형태만 간단 지원
    Cell 또는 Paragraph 객체 모두 지원
    """
    if not text:
        return
    if hasattr(cell_or_paragraph, 'paragraphs'):
        para = cell_or_paragraph.paragraphs[0]
    else:
        para = cell_or_paragraph
    parts = str(text).split("**")
    for i, part in enumerate(parts):
        run = para.add_run(part)
        if i % 2 == 1:
            run.bold = True


def add_justified_paragraph(document: Document, text: str):
    if not text:
        return
    para = document.add_paragraph()
    add_markdown_text(para, text)
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def set_table_width_by_ratio(table, ratios: list, usable_width_inches: float = 6.5):
    """
    표 컬럼 폭을 비율에 따라 고정합니다. (python-docx 버전 호환)
    
    Args:
        table: python-docx Table 객체
        ratios: 각 컬럼의 비율 리스트 (예: [0.20, 0.40, 0.40])
        usable_width_inches: 사용 가능한 페이지 폭 (인치, 기본값: 6.5 = A4 - 양쪽 1인치 마진)
    """
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    
    try:
        table.autofit = False
        
        tbl = table._tbl
        tblPr = tbl.tblPr
        if tblPr is None:
            tblPr = OxmlElement('w:tblPr')
            tbl.insert(0, tblPr)
        
        existing_layout = tblPr.find(qn('w:tblLayout'))
        if existing_layout is not None:
            tblPr.remove(existing_layout)
        
        tblLayout = OxmlElement('w:tblLayout')
        tblLayout.set(qn('w:type'), 'fixed')
        tblPr.append(tblLayout)
        
        total_width_twips = int(usable_width_inches * 1440)
        
        existing_tblW = tblPr.find(qn('w:tblW'))
        if existing_tblW is not None:
            tblPr.remove(existing_tblW)
        tblW = OxmlElement('w:tblW')
        tblW.set(qn('w:w'), str(total_width_twips))
        tblW.set(qn('w:type'), 'dxa')
        tblPr.append(tblW)
        
        for row in table.rows:
            for i, cell in enumerate(row.cells):
                if i < len(ratios):
                    cell_width = int(total_width_twips * ratios[i])
                    cell.width = cell_width
                    tc = cell._tc
                    tcPr = tc.tcPr
                    if tcPr is None:
                        tcPr = OxmlElement('w:tcPr')
                        tc.insert(0, tcPr)
                    existing_tcW = tcPr.find(qn('w:tcW'))
                    if existing_tcW is not None:
                        tcPr.remove(existing_tcW)
                    tcW = OxmlElement('w:tcW')
                    tcW.set(qn('w:w'), str(cell_width))
                    tcW.set(qn('w:type'), 'dxa')
                    tcPr.append(tcW)
        
        try:
            for i, ratio in enumerate(ratios):
                if i < len(table.columns):
                    table.columns[i].width = int(total_width_twips * ratio)
        except Exception:
            pass
            
    except Exception as e:
        print(f"[set_table_width_by_ratio] 표 폭 설정 실패 (ratios={ratios}): {e}")


def generate_satisfaction_charts(survey_data: list, total_respondents: int = 30):
    """
    만족도 조사 차트 2개 생성 (항목별 평균 점수 / 응답 분포)
    
    Returns:
        tuple: (chart1_bytes, chart2_bytes, error_log) - 에러 시 None 반환
    """
    font_prop = ensure_korean_font()
    
    try:
        if not survey_data or len(survey_data) == 0:
            return None, None, "survey_data가 비어있습니다."
        
        questions = []
        averages = []
        distributions = {'5점': [], '4점': [], '3점': [], '2점': [], '1점': []}
        
        for i, item in enumerate(survey_data):
            q_text = str(item.get('문항', f'Q{i+1}'))
            if len(q_text) > 15:
                q_text = q_text[:12] + "..."
            questions.append(q_text)
            
            s5 = int(item.get('5점', 0) or 0)
            s4 = int(item.get('4점', 0) or 0)
            s3 = int(item.get('3점', 0) or 0)
            s2 = int(item.get('2점', 0) or 0)
            s1 = int(item.get('1점', 0) or 0)
            
            total = s5 + s4 + s3 + s2 + s1
            if total > 0:
                avg = (5*s5 + 4*s4 + 3*s3 + 2*s2 + 1*s1) / total
            else:
                avg = 0
            averages.append(avg)
            
            distributions['5점'].append(s5)
            distributions['4점'].append(s4)
            distributions['3점'].append(s3)
            distributions['2점'].append(s2)
            distributions['1점'].append(s1)
        
        fig1, ax1 = plt.subplots(figsize=(6, 4))
        x = np.arange(len(questions))
        bars = ax1.bar(x, averages, color='#4472C4', edgecolor='none')
        ax1.set_ylabel('평균 점수', fontproperties=font_prop)
        ax1.set_title('항목별 평균 점수', fontproperties=font_prop, fontsize=12, fontweight='bold')
        ax1.set_xticks(x)
        ax1.set_xticklabels(questions, rotation=45, ha='right', fontproperties=font_prop, fontsize=8)
        ax1.set_ylim(0, 5)
        ax1.axhline(y=sum(averages)/len(averages) if averages else 0, color='red', linestyle='--', linewidth=1, label='전체 평균')
        ax1.legend(prop=font_prop)
        
        for bar, avg in zip(bars, averages):
            ax1.text(bar.get_x() + bar.get_width()/2, bar.get_height() + 0.1, 
                    f'{avg:.1f}', ha='center', va='bottom', fontsize=8, fontproperties=font_prop)
        
        plt.tight_layout()
        chart1_buffer = io.BytesIO()
        fig1.savefig(chart1_buffer, format='png', dpi=150, bbox_inches='tight')
        chart1_buffer.seek(0)
        plt.close(fig1)
        
        fig2, ax2 = plt.subplots(figsize=(6, 4))
        bar_width = 0.15
        colors = ['#2E7D32', '#4CAF50', '#FFC107', '#FF9800', '#F44336']
        score_labels = ['5점', '4점', '3점', '2점', '1점']
        
        for i, (score, color) in enumerate(zip(score_labels, colors)):
            offset = (i - 2) * bar_width
            ax2.bar(x + offset, distributions[score], bar_width, label=score, color=color)
        
        ax2.set_ylabel('응답 인원(명)', fontproperties=font_prop)
        ax2.set_title('응답 분포 (인원수)', fontproperties=font_prop, fontsize=12, fontweight='bold')
        ax2.set_xticks(x)
        ax2.set_xticklabels(questions, rotation=45, ha='right', fontproperties=font_prop, fontsize=8)
        ax2.legend(prop=font_prop, loc='upper right')
        
        plt.tight_layout()
        chart2_buffer = io.BytesIO()
        fig2.savefig(chart2_buffer, format='png', dpi=150, bbox_inches='tight')
        chart2_buffer.seek(0)
        plt.close(fig2)
        
        return chart1_buffer, chart2_buffer, None
        
    except Exception as e:
        plt.close('all')
        return None, None, f"차트 생성 오류: {str(e)}"


def _add_table_from_rows(doc: Document, headers, rows):
    """
    rows: list[dict] 형태를 받아서 docx 테이블로 추가
    """
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"

    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = str(h)

    for item in rows:
        r = table.add_row().cells
        for i, key in enumerate(headers.keys()):
            r[i].text = str(item.get(key, "") or "")

    doc.add_paragraph("")
    return table


# --- 리포트 생성 함수 (main.py 호출용 필수 함수들) ---
def generate_part1_report(data_dict):
    """
    PART 1 Word 문서 생성 - 전체 섹션 포함
    
    포함 섹션:
    1. 사업의 필요성 (need_1, need_2_1, need_2_2, need_2_3)
    2. 전년도 사업평가 및 환류계획 (feedback_table, total_review_table)
    3. 만족도조사 (survey_data 테이블, subjective_analysis, overall_suggestion)
    4. 사업목적 (purpose_text)
    5. 사업목표 (goals_text)
    """
    doc = Document()
    set_standard_margins(doc)
    add_left_aligned_heading(doc, "PART 1: 총괄 및 기획", 1)
    
    doc.add_heading("1. 사업의 필요성", level=2)
    
    doc.add_heading("1) 이용아동의 욕구 및 문제점", level=3)
    need_1 = data_dict.get("need_1_user_desire", "")
    if need_1:
        add_markdown_text(doc.add_paragraph(), need_1)
    
    doc.add_heading("2) 지역 환경적 특성", level=3)
    
    doc.add_heading("(1) 지역적 특성", level=4)
    need_2_1 = data_dict.get("need_2_1_regional", "")
    if need_2_1:
        add_markdown_text(doc.add_paragraph(), need_2_1)
    
    doc.add_heading("(2) 주변 환경", level=4)
    need_2_2 = data_dict.get("need_2_2_environment", "")
    if need_2_2:
        add_markdown_text(doc.add_paragraph(), need_2_2)
    
    doc.add_heading("(3) 교육적 특성", level=4)
    need_2_3 = data_dict.get("need_2_3_educational", "")
    if need_2_3:
        add_markdown_text(doc.add_paragraph(), need_2_3)
    
    doc.add_heading("2. 전년도 사업평가 및 환류계획", level=2)
    
    doc.add_heading("1) 차년도 사업 환류 계획", level=3)
    feedback_table = data_dict.get("feedback_table", [])
    if feedback_table:
        table = doc.add_table(rows=1, cols=3)
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        headers = ["영역", "문제점", "개선방안"]
        for i, h in enumerate(headers):
            hdr[i].text = h
            set_cell_background(hdr[i], "D9D9D9")
        
        for item in feedback_table:
            row = table.add_row().cells
            row[0].text = str(item.get('area', '') or '')
            add_markdown_text(row[1], str(item.get('problem', '') or ''))
            add_markdown_text(row[2], str(item.get('improvement', '') or ''))
        
        set_table_width_by_ratio(table, [0.20, 0.40, 0.40])
        doc.add_paragraph("")
    
    doc.add_heading("2) 총평", level=3)
    total_review_table = data_dict.get("total_review_table", [])
    if total_review_table:
        table = doc.add_table(rows=1, cols=2)
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        headers = ["영역", "내용"]
        for i, h in enumerate(headers):
            hdr[i].text = h
            set_cell_background(hdr[i], "D9D9D9")
        
        for item in total_review_table:
            row = table.add_row().cells
            row[0].text = str(item.get('category', '') or '')
            add_markdown_text(row[1], str(item.get('content', '') or ''))
        
        set_table_width_by_ratio(table, [0.20, 0.80])
        doc.add_paragraph("")
    
    doc.add_heading("3. 만족도조사", level=2)
    satisfaction = data_dict.get("satisfaction_survey", {})
    if satisfaction:
        total_resp = satisfaction.get('total_respondents', 0)
        if total_resp:
            doc.add_paragraph(f"총 응답 인원: {total_resp}명")
        
        survey_data = satisfaction.get('survey_data', [])
        if survey_data:
            doc.add_heading("문항별 응답 분포", level=3)
            table = doc.add_table(rows=1, cols=7)
            table.style = "Table Grid"
            hdr = table.rows[0].cells
            headers = ["문항", "5점", "4점", "3점", "2점", "1점", "평균"]
            for i, h in enumerate(headers):
                hdr[i].text = h
                set_cell_background(hdr[i], "D9D9D9")
            
            total_avg_sum = 0
            total_count = 0
            
            for item in survey_data:
                row = table.add_row().cells
                row[0].text = str(item.get('문항', '') or '')
                s5 = int(item.get('5점', 0) or 0)
                s4 = int(item.get('4점', 0) or 0)
                s3 = int(item.get('3점', 0) or 0)
                s2 = int(item.get('2점', 0) or 0)
                s1 = int(item.get('1점', 0) or 0)
                row[1].text = str(s5)
                row[2].text = str(s4)
                row[3].text = str(s3)
                row[4].text = str(s2)
                row[5].text = str(s1)
                
                total = s5 + s4 + s3 + s2 + s1
                if total > 0:
                    avg = (5*s5 + 4*s4 + 3*s3 + 2*s2 + 1*s1) / total
                    row[6].text = f"{avg:.2f}"
                    total_avg_sum += avg
                    total_count += 1
                else:
                    row[6].text = "-"
            
            score_ratio = 0.40 / 6
            set_table_width_by_ratio(table, [0.60, score_ratio, score_ratio, score_ratio, score_ratio, score_ratio, score_ratio])
            doc.add_paragraph("")
            
            if total_count > 0:
                overall_avg = total_avg_sum / total_count
                avg_para = doc.add_paragraph()
                avg_para.add_run(f"전체 평균 만족도: {overall_avg:.2f}점 (5점 만점)").bold = True
                doc.add_paragraph("")
            
            doc.add_heading("만족도 분포 차트", level=3)
            chart1, chart2, chart_error = generate_satisfaction_charts(survey_data, total_resp)
            
            if chart1 and chart2:
                chart_table = doc.add_table(rows=1, cols=2)
                chart_table.style = "Table Grid"
                chart_table.alignment = WD_TABLE_ALIGNMENT.CENTER
                
                cell1 = chart_table.rows[0].cells[0]
                cell2 = chart_table.rows[0].cells[1]
                
                para1 = cell1.paragraphs[0]
                run1 = para1.add_run()
                run1.add_picture(chart1, width=Inches(3.0))
                
                para2 = cell2.paragraphs[0]
                run2 = para2.add_run()
                run2.add_picture(chart2, width=Inches(3.0))
                
                set_table_width_by_ratio(chart_table, [0.50, 0.50])
                doc.add_paragraph("")
            elif chart_error:
                doc.add_paragraph(f"[차트 생성 실패: {chart_error}]")
        
        subjective_q = satisfaction.get('subjective_question', '')
        subjective_analysis = satisfaction.get('subjective_analysis', '')
        if subjective_analysis:
            doc.add_heading("주관식 문항 분석", level=3)
            if subjective_q:
                doc.add_paragraph(f"문항: {subjective_q}")
            add_markdown_text(doc.add_paragraph(), subjective_analysis)
        
        overall_suggestion = satisfaction.get('overall_suggestion', '')
        if overall_suggestion:
            doc.add_heading("종합 분석 및 제언", level=3)
            add_markdown_text(doc.add_paragraph(), overall_suggestion)
    
    doc.add_heading("4. 사업목적", level=2)
    purpose_text = data_dict.get("purpose_text", "")
    if purpose_text:
        add_markdown_text(doc.add_paragraph(), purpose_text)
    
    doc.add_heading("5. 사업목표", level=2)
    goals_text = data_dict.get("goals_text", "")
    if goals_text:
        add_markdown_text(doc.add_paragraph(), goals_text)
    
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def generate_part2_report(programs_dict):
    """PART 2 세부사업 Word 문서 생성 (세부사업내용 + 평가계획 테이블 포함)"""
    doc = Document()
    set_standard_margins(doc)
    add_left_aligned_heading(doc, "PART 2: 세부 사업 계획", 1)

    # 5대영역 순서 정의
    category_order = ["보호", "교육", "문화", "정서지원", "지역사회연계"]
    first_category = True

    if isinstance(programs_dict, dict):
        # 정렬된 순서로 영역 처리
        sorted_categories = sorted(
            [c for c in programs_dict.keys() if not c.startswith('_')],
            key=lambda x: category_order.index(x) if x in category_order else 999
        )
        
        for category in sorted_categories:
            payload = programs_dict[category]
            
            # 페이지 나눔 (첫 영역 제외)
            if not first_category:
                doc.add_page_break()
            first_category = False
            
            doc.add_heading(str(category), level=2)
            
            detail_table = payload.get('detail_table', []) if isinstance(payload, dict) else []
            
            # 기대효과 100자 제한 적용 (이중 안전장치 - 워드 출력 직전)
            effect_map = {}  # program_name -> expected_effect 매핑 (평가계획 연동용)
            if detail_table:
                for item in detail_table:
                    original_effect = str(item.get('expected_effect', '') or '')
                    truncated_effect = truncate_expected_effect(original_effect)
                    item['expected_effect'] = truncated_effect
                    
                    # 평가계획 연동을 위한 매핑 저장
                    prog_name = str(item.get('program_name', '') or '')
                    if prog_name:
                        effect_map[prog_name] = truncated_effect
            
            if detail_table:
                doc.add_heading("세부사업내용", level=3)
                detail_headers = ["세부영역", "프로그램명", "기대효과", "대상", "인원", "주기", "계획내용"]
                table = doc.add_table(rows=1, cols=len(detail_headers))
                table.style = "Table Grid"
                
                hdr = table.rows[0].cells
                for i, h in enumerate(detail_headers):
                    hdr[i].text = h
                    set_cell_background(hdr[i], "D9D9D9")
                
                for item in detail_table:
                    row = table.add_row().cells
                    row[0].text = str(item.get('sub_area', '') or '')
                    row[1].text = str(item.get('program_name', '') or '')
                    row[2].text = str(item.get('expected_effect', '') or '')
                    row[3].text = str(item.get('target', '') or '')
                    row[4].text = str(item.get('count', '') or '')
                    row[5].text = str(item.get('cycle', '') or '')
                    content = str(item.get('content', '') or '')
                    add_markdown_text(row[6], content)
                
                set_table_width_by_ratio(table, [0.08, 0.08, 0.30, 0.08, 0.08, 0.08, 0.30])
                doc.add_paragraph("")
            
            eval_table = payload.get('eval_table', []) if isinstance(payload, dict) else []
            
            # 기존 스키마 호환성 + 기대효과 연동
            for item in eval_table:
                # 기존 스키마 호환: eval_tool/eval_timing → main_plan/eval_method
                if 'eval_tool' in item and 'main_plan' not in item:
                    item['main_plan'] = item.pop('eval_tool', '')
                if 'eval_timing' in item and 'eval_method' not in item:
                    item['eval_method'] = item.pop('eval_timing', '')
                if 'sub_area' not in item:
                    item['sub_area'] = ''
                
                # 기대효과가 없으면 detail_table에서 연동
                prog_name = str(item.get('program_name', '') or '')
                if not item.get('expected_effect') and prog_name in effect_map:
                    item['expected_effect'] = effect_map[prog_name]
            
            if eval_table:
                doc.add_heading("평가계획", level=3)
                eval_headers = ["세부영역", "프로그램명", "기대효과", "평가계획", "평가방법"]
                table = doc.add_table(rows=1, cols=len(eval_headers))
                table.style = "Table Grid"
                
                hdr = table.rows[0].cells
                for i, h in enumerate(eval_headers):
                    hdr[i].text = h
                    set_cell_background(hdr[i], "D9D9D9")
                
                for item in eval_table:
                    row = table.add_row().cells
                    row[0].text = str(item.get('sub_area', '') or '')
                    prog_name = str(item.get('program_name', '') or '')
                    row[1].text = prog_name
                    
                    # 기대효과 연동: detail_table의 값을 사용 (없으면 eval_table 값 사용 후 100자 제한)
                    if prog_name in effect_map:
                        effect = effect_map[prog_name]
                    else:
                        effect = truncate_expected_effect(str(item.get('expected_effect', '') or ''))
                    add_markdown_text(row[2], effect)
                    row[3].text = str(item.get('main_plan', '') or '')
                    row[4].text = str(item.get('eval_method', '') or '')
                
                # 평가계획 표 비율: 15/15/40/15/15
                set_table_width_by_ratio(table, [0.15, 0.15, 0.40, 0.15, 0.15])
                doc.add_paragraph("")
            
            if not detail_table and not eval_table:
                doc.add_paragraph("등록된 사업이 없습니다.")
            
            doc.add_paragraph("")

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def generate_monthly_report(monthly_list, period):
    """
    (유지) 단일 기간용 월별 계획 템플릿
    """
    doc = Document()
    set_standard_margins(doc)
    add_left_aligned_heading(doc, f"{period} 월별 계획", 1)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def generate_monthly_program_report(monthly_plan_dict,
                                    months_list,
                                    title=None):
    """
    PART 3 상반기 월별 사업계획 Word 문서 생성
    
    months_list 예: ["1월","2월","3월","4월","5월","6월"]
    monthly_plan_dict 구조(예시):
      {
        "1월": [{"big_category":..., "mid_category":..., "program_name":..., "target":..., "staff":..., "content":...}, ...],
        ...
      }
    """
    doc = Document()
    set_standard_margins(doc)

    if title is None:
        if months_list:
            title = f"PART 3: {months_list[0]}~{months_list[-1]} 월별 사업계획"
        else:
            title = "PART 3: 월별 사업계획"

    add_left_aligned_heading(doc, title, 1)

    headers = {
        "big_category": "대분류",
        "mid_category": "중분류",
        "program_name": "프로그램명",
        "target": "참여자",
        "staff": "수행인력",
        "content": "사업내용",
    }

    for month in months_list:
        doc.add_heading(month, level=2)

        rows = monthly_plan_dict.get(month, []) if isinstance(
            monthly_plan_dict, dict) else []
        rows = rows or []

        if not rows:
            doc.add_paragraph("등록된 사업이 없습니다.")
            doc.add_paragraph("")
            continue

        table = doc.add_table(rows=1, cols=len(headers))
        table.style = "Table Grid"

        hdr_cells = table.rows[0].cells
        for i, h in enumerate(headers.values()):
            hdr_cells[i].text = str(h)
            set_cell_background(hdr_cells[i], "D9D9D9")

        for item in rows:
            r = table.add_row().cells
            r[0].text = str(item.get("big_category", "") or "")
            r[1].text = str(item.get("mid_category", "") or "")
            r[2].text = str(item.get("program_name", "") or "")
            r[3].text = str(item.get("target", "") or "")
            r[4].text = str(item.get("staff", "") or "")
            content = str(item.get("content", "") or "")
            add_markdown_text(r[5], content)

        doc.add_paragraph("")

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def generate_budget_evaluation_report(budget_eval_data):
    """예산 및 평가 Word 문서 생성"""
    doc = Document()
    set_standard_margins(doc)
    add_left_aligned_heading(doc, "PART 4: 예산 및 평가", 1)

    budget_table = budget_eval_data.get('budget_table', []) if isinstance(budget_eval_data, dict) else []
    if budget_table:
        doc.add_heading("예산계획", level=2)
        headers = ["항목", "금액", "세부내용"]
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = "Table Grid"
        
        hdr = table.rows[0].cells
        for i, h in enumerate(headers):
            hdr[i].text = h
            set_cell_background(hdr[i], "D9D9D9")
        
        for item in budget_table:
            row = table.add_row().cells
            row[0].text = str(item.get('category', '') or '')
            row[1].text = str(item.get('amount', '') or '')
            row[2].text = str(item.get('details', '') or '')
        
        doc.add_paragraph("")
    
    feedback_summary = budget_eval_data.get('feedback_summary', []) if isinstance(budget_eval_data, dict) else []
    if feedback_summary:
        doc.add_heading("환류 요약", level=2)
        headers = ["영역", "문제점", "개선계획"]
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = "Table Grid"
        
        hdr = table.rows[0].cells
        for i, h in enumerate(headers):
            hdr[i].text = h
            set_cell_background(hdr[i], "D9D9D9")
        
        for item in feedback_summary:
            row = table.add_row().cells
            row[0].text = str(item.get('area', '') or '')
            problem = str(item.get('problem', '') or '')
            add_markdown_text(row[1], problem)
            plan = str(item.get('plan', '') or '')
            add_markdown_text(row[2], plan)
        
        doc.add_paragraph("")

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def generate_part4_full_report(monthly_plan_dict, months_list, budget_eval_data):
    """PART 4 전체 Word 문서 생성 (하반기 월별계획 + 예산 + 환류요약)"""
    doc = Document()
    set_standard_margins(doc)
    add_left_aligned_heading(doc, "PART 4: 하반기 월별 사업계획 및 평가", 1)
    
    doc.add_heading("하반기 월별 사업계획 (7월~12월)", level=2)
    
    headers = {
        "big_category": "대분류",
        "mid_category": "중분류",
        "program_name": "프로그램명",
        "target": "참여자",
        "staff": "수행인력",
        "content": "사업내용",
    }
    
    for month in months_list:
        doc.add_heading(month, level=3)
        
        rows = monthly_plan_dict.get(month, []) if isinstance(monthly_plan_dict, dict) else []
        rows = rows or []
        
        if not rows:
            doc.add_paragraph("등록된 사업이 없습니다.")
            continue
        
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = "Table Grid"
        
        hdr_cells = table.rows[0].cells
        for i, h in enumerate(headers.values()):
            hdr_cells[i].text = str(h)
            set_cell_background(hdr_cells[i], "D9D9D9")
        
        for item in rows:
            r = table.add_row().cells
            r[0].text = str(item.get("big_category", "") or "")
            r[1].text = str(item.get("mid_category", "") or "")
            r[2].text = str(item.get("program_name", "") or "")
            r[3].text = str(item.get("target", "") or "")
            r[4].text = str(item.get("staff", "") or "")
            content = str(item.get("content", "") or "")
            add_markdown_text(r[5], content)
        
        doc.add_paragraph("")
    
    budget_table = budget_eval_data.get('budget_table', []) if isinstance(budget_eval_data, dict) else []
    if budget_table:
        doc.add_heading("예산계획", level=2)
        headers = ["항목", "금액", "세부내용"]
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = "Table Grid"
        
        hdr = table.rows[0].cells
        for i, h in enumerate(headers):
            hdr[i].text = h
            set_cell_background(hdr[i], "D9D9D9")
        
        for item in budget_table:
            row = table.add_row().cells
            row[0].text = str(item.get('category', '') or '')
            row[1].text = str(item.get('amount', '') or '')
            row[2].text = str(item.get('details', '') or '')
        
        doc.add_paragraph("")
    
    feedback_summary = budget_eval_data.get('feedback_summary', []) if isinstance(budget_eval_data, dict) else []
    if feedback_summary:
        doc.add_heading("환류 요약", level=2)
        headers = ["영역", "문제점", "개선계획"]
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = "Table Grid"
        
        hdr = table.rows[0].cells
        for i, h in enumerate(headers):
            hdr[i].text = h
            set_cell_background(hdr[i], "D9D9D9")
        
        for item in feedback_summary:
            row = table.add_row().cells
            row[0].text = str(item.get('area', '') or '')
            problem = str(item.get('problem', '') or '')
            add_markdown_text(row[1], problem)
            plan = str(item.get('plan', '') or '')
            add_markdown_text(row[2], plan)
    
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def generate_full_report(data_dict):
    doc = Document()
    set_standard_margins(doc)
    doc.add_heading("2025 연간 사업 계획서", 0).alignment = WD_ALIGN_PARAGRAPH.CENTER

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
