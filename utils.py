import os
import json
import re
import io
import ast
import streamlit as st
import google.generativeai as genai

GEMINI_API_KEY = os.environ.get("GEMINI_API_KEY")
GEMINI_MODEL = os.environ.get("GEMINI_MODEL", "gemini-2.0-flash")

if GEMINI_API_KEY:
    genai.configure(api_key=GEMINI_API_KEY)

DEFAULT_GEN_CONFIG = genai.GenerationConfig(
    response_mime_type="application/json",
    temperature=0.5,
    max_output_tokens=8192
)

JSON_PROMPT_RULES = """[필수 출력 규칙]
- 반드시 단일 JSON 객체만 출력 (배열 금지)
- 마크다운/코드펜스(```) 절대 금지
- 문자열 내 줄바꿈은 반드시 \\n으로 표현 (실제 개행 금지)
- 키 이름은 지정된 스키마 그대로 사용
- 마지막은 반드시 }로 끝나야 함
- 각 필드는 600자 이내로 요약, 최대 5개 불릿만 사용
"""

LABEL_PATTERNS = [
    (r'(?:프로그램명|프로그램\s*명칭|사업명)[:\s]*([^\n]{2,100})', 'program_name'),
    (r'(?:일자|일시|실시\s*일)[:\s]*([^\n]{2,50})', 'date'),
    (r'(?:담당자|담당|수행인력)[:\s]*([^\n]{2,30})', 'staff'),
    (r'(?:대상자|대상|참여자|참가자)[:\s]*([^\n]{2,50})', 'target'),
    (r'(?:목적|사업\s*목적)[:\s]*([^\n]{2,300})', 'purpose'),
    (r'(?:목표|사업\s*목표)[:\s]*([^\n]{2,300})', 'goal'),
    (r'(?:기대\s*효과|효과)[:\s]*([^\n]{2,200})', 'effect'),
    (r'(?:평가|평가\s*내용|평가\s*결과)[:\s]*([^\n]{2,300})', 'evaluation'),
    (r'(?:향후\s*계획|차년도\s*계획|개선\s*계획)[:\s]*([^\n]{2,300})', 'plan'),
    (r'(?:주기|횟수|빈도)[:\s]*([^\n]{2,30})', 'cycle'),
    (r'(?:영역|분류|카테고리)[:\s]*([^\n]{2,30})', 'category'),
]

MONTH_PATTERNS = [
    (r'(\d{1,2})월', 'month_num'),
    (r'(1|2|3|4|5|6|7|8|9|10|11|12)\s*월', 'month_num'),
    (r'(일|이|삼|사|오|육|칠|팔|구|십|십일|십이)\s*월', 'month_korean'),
]

KOREAN_MONTH_MAP = {
    '일': 1, '이': 2, '삼': 3, '사': 4, '오': 5, '육': 6,
    '칠': 7, '팔': 8, '구': 9, '십': 10, '십일': 11, '십이': 12
}

CYCLE_TO_MONTHS = {
    '연중': list(range(1, 13)),
    '매월': list(range(1, 13)),
    '주1회': list(range(1, 13)),
    '주2회': list(range(1, 13)),
    '주3회': list(range(1, 13)),
    '주4회': list(range(1, 13)),
    '주5회': list(range(1, 13)),
    '상시': list(range(1, 13)),
    '여름방학': [7, 8],
    '겨울방학': [1, 2, 12],
    '방학': [1, 2, 7, 8, 12],
    '분기': [3, 6, 9, 12],
    '분기1회': [3, 6, 9, 12],
    '반기': [6, 12],
    '연2회': [1, 7],
    '연1회': [6],
}


def load_guideline_rules() -> dict:
    """guidelines_template.json에서 작성지침 규칙을 로드합니다."""
    import os
    result = {"_load_status": "unknown", "_load_error": None}
    
    if not os.path.exists('guidelines_template.json'):
        result["_load_status"] = "file_not_found"
        result["_load_error"] = "guidelines_template.json 파일이 존재하지 않습니다."
        print(f"[load_guideline_rules] {result['_load_error']}")
        return result
    
    try:
        with open('guidelines_template.json', 'r', encoding='utf-8') as f:
            data = json.load(f)
            data["_load_status"] = "success"
            data["_load_error"] = None
            print(f"[load_guideline_rules] 성공. part1 키: {list(data.get('part1', {}).keys())}")
            return data
    except json.JSONDecodeError as e:
        result["_load_status"] = "json_error"
        result["_load_error"] = f"JSON 파싱 오류: {str(e)}"
        print(f"[load_guideline_rules] {result['_load_error']}")
        return result
    except Exception as e:
        result["_load_status"] = "error"
        result["_load_error"] = f"알 수 없는 오류: {str(e)}"
        print(f"[load_guideline_rules] {result['_load_error']}")
        return result


def count_chars_no_space(text: str) -> int:
    """공백을 제외한 글자수를 계산합니다."""
    if not text:
        return 0
    return len(text.replace(' ', '').replace('\n', '').replace('\t', ''))


def smart_truncate(text: str, max_chars_no_space: int) -> str:
    """문장 단위로 스마트하게 자릅니다 (마침표 기준)."""
    if count_chars_no_space(text) <= max_chars_no_space:
        return text
    
    sentences = re.split(r'(?<=[.!?。])\s*', text)
    result = ""
    for sentence in sentences:
        test = result + sentence
        if count_chars_no_space(test) > max_chars_no_space:
            break
        result = test + " "
    
    if not result.strip():
        chars_counted = 0
        cut_idx = 0
        for i, ch in enumerate(text):
            if ch not in ' \n\t':
                chars_counted += 1
            if chars_counted >= max_chars_no_space:
                cut_idx = i
                break
        result = text[:cut_idx+1]
    
    return result.strip()


def validate_and_fix_text(text: str, max_chars_no_space: int, field_name: str = "", max_retries: int = 2) -> str:
    """
    텍스트 글자수 검사 및 자동 보정.
    1) 공백 제외 글자수 검사
    2) 초과 시 Gemini에 재작성 요청 (최대 2회)
    3) 그래도 초과면 스마트 컷
    """
    if not text:
        return text
    
    current_count = count_chars_no_space(text)
    
    if current_count <= max_chars_no_space:
        return text
    
    if not GEMINI_API_KEY:
        return smart_truncate(text, max_chars_no_space)
    
    result_text = text
    
    for attempt in range(max_retries):
        current_count = count_chars_no_space(result_text)
        if current_count <= max_chars_no_space:
            return result_text
        
        try:
            rewrite_prompt = f"""다음 텍스트를 공백 제외 {max_chars_no_space}자 이내로 줄여주세요.
현재 글자수: {current_count}자 (공백 제외)
목표 글자수: {max_chars_no_space}자 이내 (공백 제외)

핵심 내용은 유지하면서 간결하게 요약해주세요.
● 불릿 형식은 유지하세요.
마크다운이나 설명 없이 수정된 텍스트만 출력하세요.

원문:
{result_text}"""
            
            model = genai.GenerativeModel(
                model_name=GEMINI_MODEL,
                generation_config=genai.GenerationConfig(temperature=0.3, max_output_tokens=2048)
            )
            response = model.generate_content(rewrite_prompt)
            if response.text:
                result_text = response.text.strip()
                print(f"[validate_text] Retry {attempt+1}: {current_count} -> {count_chars_no_space(result_text)} chars")
        except Exception as e:
            print(f"[validate_text] Gemini rewrite error: {e}")
            break
    
    if count_chars_no_space(result_text) > max_chars_no_space:
        result_text = smart_truncate(result_text, max_chars_no_space)
        print(f"[validate_text] Smart truncated to {count_chars_no_space(result_text)} chars")
    
    return result_text


def extract_months_from_text(text: str) -> list:
    """텍스트에서 월 정보를 추출합니다."""
    months = set()
    
    for pattern, ptype in MONTH_PATTERNS:
        matches = re.findall(pattern, text)
        for match in matches:
            if ptype == 'month_num':
                try:
                    m = int(match)
                    if 1 <= m <= 12:
                        months.add(m)
                except ValueError:
                    pass
            elif ptype == 'month_korean':
                m = KOREAN_MONTH_MAP.get(match)
                if m:
                    months.add(m)
    
    return sorted(list(months))


def extract_months_from_cycle(cycle: str) -> list:
    """주기 문자열에서 해당 월 목록을 추출합니다."""
    if not cycle:
        return list(range(1, 13))
    
    cycle_clean = cycle.strip().replace(' ', '')
    
    for key, months in CYCLE_TO_MONTHS.items():
        if key in cycle_clean:
            return months
    
    extracted = extract_months_from_text(cycle)
    if extracted:
        return extracted
    
    range_match = re.search(r'(\d{1,2})월?\s*[~\-]\s*(\d{1,2})월?', cycle)
    if range_match:
        start, end = int(range_match.group(1)), int(range_match.group(2))
        if 1 <= start <= 12 and 1 <= end <= 12:
            if start <= end:
                return list(range(start, end + 1))
            else:
                return list(range(start, 13)) + list(range(1, end + 1))
    
    return list(range(1, 13))


def bucket_programs_by_month(file_summaries: list) -> dict:
    """파일 요약에서 프로그램을 월별로 분류합니다."""
    month_bucket = {m: [] for m in range(1, 13)}
    
    for summary in file_summaries:
        labels = summary.get('labels', {})
        program_names = labels.get('program_name', [])
        dates = labels.get('date', [])
        cycles = labels.get('cycle', [])
        targets = labels.get('target', [])
        staffs = labels.get('staff', [])
        
        for i, prog_name in enumerate(program_names):
            date_str = dates[i] if i < len(dates) else ""
            cycle_str = cycles[i] if i < len(cycles) else ""
            target_str = targets[i] if i < len(targets) else "전체아동"
            staff_str = staffs[i] if i < len(staffs) else "돌봄교사"
            
            months_from_date = extract_months_from_text(date_str)
            months_from_cycle = extract_months_from_cycle(cycle_str)
            
            if months_from_date:
                target_months = months_from_date
            else:
                target_months = months_from_cycle
            
            program_info = {
                "program_name": prog_name,
                "target": target_str,
                "staff": staff_str,
                "cycle": cycle_str,
                "date": date_str,
                "source_file": summary.get('filename', '')
            }
            
            for m in target_months:
                month_bucket[m].append(program_info)
    
    return month_bucket


def get_default_monthly_template() -> dict:
    """빈 월에 채울 정기운영 템플릿을 반환합니다."""
    return {
        "big_category": "보호",
        "mid_category": "생활",
        "program_name": "일상생활지도",
        "target": "전체아동",
        "staff": "돌봄교사",
        "content": "● **정기운영**: 일상생활 및 위생관리 지도"
    }


def extract_labels_from_text(text: str) -> dict:
    """규칙 기반으로 텍스트에서 라벨별 값을 추출 (정규식 활용)."""
    result = {}
    for pattern, key in LABEL_PATTERNS:
        matches = re.findall(pattern, text, re.IGNORECASE | re.MULTILINE)
        if matches:
            result[key] = [m.strip() for m in matches if m.strip()]
    return result


def extract_file_summaries(uploaded_files: list) -> list:
    """각 업로드 파일에서 규칙 기반 라벨 추출하여 요약 dict 목록 반환."""
    summaries = []
    for uf in uploaded_files:
        file_text = read_uploaded_file(uf)
        uf.seek(0)
        
        if not file_text:
            continue
        
        labels = extract_labels_from_text(file_text)
        
        summary = {
            "filename": uf.name,
            "text_length": len(file_text),
            "labels": labels,
            "text_preview": file_text[:500] if len(file_text) > 500 else file_text
        }
        summaries.append(summary)
    
    return summaries


def summaries_to_compact_text(summaries: list) -> str:
    """파일 요약 목록을 Gemini 입력용 간결한 텍스트로 변환."""
    lines = []
    for i, s in enumerate(summaries, 1):
        lines.append(f"[파일 {i}: {s['filename']}]")
        for key, values in s.get('labels', {}).items():
            if values:
                lines.append(f"  {key}: {', '.join(values[:3])}")
        lines.append(f"  미리보기: {s.get('text_preview', '')[:200]}...")
        lines.append("")
    return "\n".join(lines)


def read_pdf(file) -> str:
    """Extract text from PDF file using pdfplumber."""
    try:
        import pdfplumber
        text_content = []
        with pdfplumber.open(file) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text_content.append(page_text)
        return "\n".join(text_content)
    except Exception as e:
        st.error(f"PDF 파일 읽기 오류: {str(e)}")
        return ""


def read_docx(file) -> str:
    """Extract text from DOCX file."""
    try:
        from docx import Document
        doc = Document(file)
        text_content = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_content.append(paragraph.text)
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells]
                text_content.append(" | ".join(row_text))
        return "\n".join(text_content)
    except Exception as e:
        st.error(f"DOCX 파일 읽기 오류: {str(e)}")
        return ""


def read_uploaded_file(uploaded_file) -> str:
    """Read content from various file types."""
    file_type = uploaded_file.name.split('.')[-1].lower()
    
    if file_type == 'pdf':
        return read_pdf(uploaded_file)
    elif file_type == 'docx':
        return read_docx(uploaded_file)
    elif file_type in ['txt', 'csv']:
        return uploaded_file.read().decode('utf-8')
    elif file_type == 'hwp':
        st.warning("HWP 파일은 현재 텍스트 추출이 제한됩니다. 가능하면 PDF나 DOCX로 변환해주세요.")
        return ""
    else:
        st.error(f"지원하지 않는 파일 형식입니다: {file_type}")
        return ""


def process_multiple_files(uploaded_files: list) -> str:
    """Process multiple uploaded files and combine their text content."""
    all_text = ""
    
    for uploaded_file in uploaded_files:
        file_text = read_uploaded_file(uploaded_file)
        if file_text:
            all_text += f"\n--- {uploaded_file.name} 시작 ---\n"
            all_text += file_text
            all_text += f"\n--- {uploaded_file.name} 끝 ---\n"
        uploaded_file.seek(0)
    
    return all_text.strip()


def get_api_key():
    """Get Gemini API key from environment variables."""
    api_key = os.environ.get('GEMINI_API_KEY')
    if not api_key:
        st.error("GEMINI_API_KEY가 설정되지 않았습니다. Replit Secrets에서 GEMINI_API_KEY를 추가해주세요.")
        return None
    return api_key


def _remove_code_blocks(text: str) -> str:
    """Remove markdown code blocks from text."""
    text = re.sub(r'^```json\s*\n?', '', text.strip(), flags=re.MULTILINE)
    text = re.sub(r'^```\s*\n?', '', text, flags=re.MULTILINE)
    text = re.sub(r'\n?```\s*$', '', text, flags=re.MULTILINE)
    text = text.replace('```json', '').replace('```', '')
    return text.strip()


def _extract_balanced_json(text: str, start_char: str) -> str:
    """Extract balanced JSON starting from start_char ({ or [)."""
    if start_char == '{':
        open_char, close_char = '{', '}'
    else:
        open_char, close_char = '[', ']'
    
    start_idx = text.find(start_char)
    if start_idx == -1:
        return None
    
    depth = 0
    in_string = False
    escape_next = False
    
    for i, ch in enumerate(text[start_idx:], start=start_idx):
        if escape_next:
            escape_next = False
            continue
        if ch == '\\':
            escape_next = True
            continue
        if ch == '"' and not escape_next:
            in_string = not in_string
            continue
        if in_string:
            continue
        if ch == open_char:
            depth += 1
        elif ch == close_char:
            depth -= 1
            if depth == 0:
                return text[start_idx:i+1]
    
    return text[start_idx:]


def _extract_json_from_text(raw: str) -> dict:
    """Extract and parse JSON from raw text with robust handling."""
    if not raw or not raw.strip():
        return None
    
    text = _remove_code_blocks(raw.strip())
    
    first_brace = text.find('{')
    first_bracket = text.find('[')
    
    if first_brace == -1 and first_bracket == -1:
        return None
    
    if first_brace != -1:
        json_str = _extract_balanced_json(text, '{')
        if json_str:
            try:
                return json.loads(json_str)
            except json.JSONDecodeError:
                pass
            try:
                return ast.literal_eval(json_str)
            except:
                pass
    
    if first_bracket != -1:
        json_str = _extract_balanced_json(text, '[')
        if json_str:
            try:
                return json.loads(json_str)
            except json.JSONDecodeError:
                pass
            try:
                return ast.literal_eval(json_str)
            except:
                pass
    
    json_patterns = [r'\{[\s\S]*\}', r'\[[\s\S]*\]']
    for pattern in json_patterns:
        matches = re.findall(pattern, text)
        if matches:
            matches.sort(key=len, reverse=True)
            for candidate in matches:
                try:
                    return json.loads(candidate)
                except json.JSONDecodeError:
                    continue
    
    return None


def _ensure_dict(parsed) -> dict:
    """Ensure parsed result is a dict, unwrapping single-item lists."""
    if parsed is None:
        return None
    if isinstance(parsed, dict):
        return parsed
    if isinstance(parsed, list):
        if len(parsed) == 1 and isinstance(parsed[0], dict):
            return parsed[0]
        return None
    return None


def normalize_table_rows(rows, columns: list, fill_value: str = "") -> list:
    """Normalize table rows to ensure all columns exist."""
    if not isinstance(rows, list):
        return []
    
    normalized = []
    for row in rows:
        if not isinstance(row, dict):
            continue
        normalized_row = {}
        for col in columns:
            normalized_row[col] = row.get(col, fill_value)
        normalized.append(normalized_row)
    
    return normalized


def _is_truncated(raw_text: str) -> bool:
    """응답이 중간에 끊겼는지 확인 (닫는 괄호 없음)."""
    if not raw_text or not raw_text.strip():
        return True
    text = raw_text.strip()
    if not text.endswith('}') and not text.endswith(']'):
        return True
    open_braces = text.count('{') - text.count('}')
    open_brackets = text.count('[') - text.count(']')
    if open_braces > 0 or open_brackets > 0:
        return True
    return False


def safe_gemini_json(prompt: str, system_instruction: str = None, max_retries: int = 2, shorter_on_retry: bool = True) -> dict:
    """
    Safe Gemini JSON generation with automatic retry on parsing failure.
    Returns a dict or raises an exception with raw text for debugging.
    If shorter_on_retry=True, retry prompts will request more concise output.
    """
    if not GEMINI_API_KEY:
        raise ValueError("GEMINI_API_KEY가 설정되지 않았습니다.")
    
    print(f"[Gemini] model={GEMINI_MODEL}")
    
    full_prompt = f"{JSON_PROMPT_RULES}\n\n{prompt}"
    
    model = genai.GenerativeModel(
        model_name=GEMINI_MODEL,
        system_instruction=system_instruction,
        generation_config=DEFAULT_GEN_CONFIG
    )
    
    last_raw_text = ""
    truncated = False
    
    for attempt in range(max_retries + 1):
        try:
            if attempt > 0:
                shorter_hint = ""
                if shorter_on_retry and truncated:
                    shorter_hint = """
[중요] 이전 응답이 중간에 끊겼습니다. 더 짧게 요약하세요:
- 각 필드 400자 이내
- 불릿 포인트 최대 3개
- 테이블 행 최대 5개"""
                
                retry_instruction = f"""이전 응답이 유효한 JSON이 아니었습니다.
{shorter_hint}
반드시:
1) 오직 JSON 객체 1개만 출력
2) 모든 문자열 줄바꿈은 \\n으로
3) 마크다운/설명문/코드펜스 금지
4) 첫 문자 '{{'로 시작, 마지막 '}}'로 끝

{prompt}"""
                response = model.generate_content(retry_instruction)
                print(f"[Gemini retry#{attempt}] truncated={truncated}")
            else:
                response = model.generate_content(full_prompt)
            
            raw_text = response.text if response.text else ""
            last_raw_text = raw_text
            
            preview = raw_text[:400] if raw_text else "(EMPTY)"
            tail = raw_text[-200:] if len(raw_text) > 200 else raw_text
            print(f"[Gemini raw preview] {preview}")
            print(f"[Gemini raw tail] ...{tail}")
            print(f"[Gemini raw length] {len(raw_text)}")
            
            truncated = _is_truncated(raw_text)
            if truncated:
                print(f"[Gemini] 응답 끊김 감지, 재시도...")
                continue
            
            if not raw_text.strip():
                continue
            
            try:
                parsed = json.loads(raw_text)
                result = _ensure_dict(parsed)
                if result is not None:
                    return result
            except json.JSONDecodeError:
                pass
            
            parsed = _extract_json_from_text(raw_text)
            result = _ensure_dict(parsed)
            if result is not None:
                return result
                
        except Exception as e:
            print(f"[Gemini error] {str(e)}")
            last_raw_text = str(e)
    
    raise ValueError(f"JSON 파싱 최종 실패. 원문:\n{last_raw_text[:2000]}")


def generate_part1(compact_text: str) -> dict:
    """Part1 총괄/기획 영역만 생성."""
    system_instruction = """당신은 연간 사업계획서 작성 전문가입니다.
한국어로 작성하고, 이모지 사용 금지, ● 기호만 사용하세요.
출력: 오직 JSON 객체 1개만. 마크다운/코드펜스 금지."""

    prompt = f"""다음 파일 요약을 기반으로 Part1 (총괄/기획) JSON을 생성하세요.

{compact_text}

출력 스키마:
{{
  "need_1_user_desire": "이용아동 욕구 (400자, ●불릿 3개)",
  "need_2_1_regional": "지역적 특성 (400자)",
  "need_2_2_environment": "주변환경 (400자)",
  "need_2_3_educational": "교육적 특성 (400자)",
  "feedback_table": [
    {{"area": "보호/교육/문화/정서지원/지역사회연계", "problem": "문제점 3개", "improvement": "개선방안 3개"}}
  ],
  "total_review_table": [
    {{"category": "운영평가/아동평가/프로그램평가/후원활동측면/환류방안", "content": "상세 내용"}}
  ],
  "satisfaction_survey": {{
    "total_respondents": 30,
    "survey_data": [{{"문항": "1. 질문", "5점": 15, "4점": 10, "3점": 3, "2점": 1, "1점": 1}}],
    "subjective_analysis": "주관식 분석 (300자)",
    "overall_suggestion": "종합 제언 (300자)"
  }},
  "purpose_text": "사업목적 (300자)",
  "goals_text": "사업목표 ●불릿 5개 (각 200자)"
}}

[중요] 각 필드 600자 이내, 테이블 각 5행 이내로 제한."""
    
    try:
        return safe_gemini_json(prompt, system_instruction, max_retries=2)
    except ValueError:
        return None


def generate_part2(compact_text: str) -> dict:
    """Part2 세부사업 영역만 생성."""
    system_instruction = """당신은 지역아동센터 프로그램 기획 전문가입니다.
한국어로 작성, 이모지 금지, ● 기호만 사용.
출력: 오직 JSON 객체 1개만."""

    prompt = f"""다음 파일 요약을 기반으로 Part2 (세부사업) JSON을 생성하세요.

{compact_text}

출력 스키마 (5개 카테고리):
{{
  "보호": {{
    "subcategories": ["생활", "안전", "가족기능강화"],
    "detail_table": [
      {{"sub_area": "생활", "program_name": "프로그램명", "expected_effect": "기대효과", "target": "대상", "count": "인원", "cycle": "주기", "content": "● 내용"}}
    ],
    "eval_table": [
      {{"sub_area": "생활", "program_name": "프로그램명", "expected_effect": "● 효과", "main_plan": "주요계획", "eval_method": "평가방법"}}
    ]
  }},
  "교육": {{"subcategories": ["성장과권리", "학습", "특기적성"], "detail_table": [...], "eval_table": [...]}},
  "문화": {{"subcategories": ["체험활동"], "detail_table": [...], "eval_table": [...]}},
  "정서지원": {{"subcategories": ["상담"], "detail_table": [...], "eval_table": [...]}},
  "지역사회연계": {{"subcategories": ["연계"], "detail_table": [...], "eval_table": [...]}}
}}

[중요] 각 카테고리당 detail_table 3행, eval_table 3행 이내."""
    
    try:
        return safe_gemini_json(prompt, system_instruction, max_retries=2)
    except ValueError:
        return None


def generate_part3(compact_text: str, month_bucket: dict = None, guideline_rules: dict = None) -> dict:
    """Part3 상반기 월별계획 (1~6월) 생성. month_bucket이 있으면 해당 월에 배치."""
    
    rules = guideline_rules.get('part3', {}).get('monthly_program', {}) if guideline_rules else {}
    max_programs = rules.get('max_programs_per_month', 8)
    content_max = rules.get('columns', {}).get('content', {}).get('max_chars_no_space', 200)
    
    pre_bucketed = {}
    if month_bucket:
        for m in range(1, 7):
            month_key = f"{m}월"
            programs = month_bucket.get(m, [])[:max_programs]
            pre_bucketed[month_key] = [
                {
                    "program_name": p.get('program_name', ''),
                    "target": p.get('target', '전체아동'),
                    "staff": p.get('staff', '돌봄교사'),
                    "cycle": p.get('cycle', '')
                } for p in programs
            ]
    
    pre_bucket_info = ""
    if pre_bucketed:
        pre_bucket_info = "\n[월별 프로그램 배치 정보 - 반드시 이 월에 해당 프로그램을 배치하세요]:\n"
        for month, progs in pre_bucketed.items():
            if progs:
                prog_names = ", ".join([p['program_name'] for p in progs])
                pre_bucket_info += f"  {month}: {prog_names}\n"
    
    system_instruction = """당신은 사업계획 일정 전문가입니다.
한국어, 이모지 금지, ● 기호만 사용.
출력: 오직 JSON 객체 1개만."""

    prompt = f"""다음 파일 요약을 기반으로 Part3 (상반기 1~6월) 월별계획 JSON을 생성하세요.

{compact_text}
{pre_bucket_info}

출력 스키마:
{{
  "1월": [
    {{"big_category": "보호", "mid_category": "생활", "program_name": "급식관리", "target": "전체아동", "staff": "돌봄교사", "content": "● 내용"}}
  ],
  "2월": [...],
  "3월": [...],
  "4월": [...],
  "5월": [...],
  "6월": [...]
}}

[중요] 
- 각 월당 프로그램 {max_programs}개 이내
- 각 content {content_max}자 (공백 제외) 이내
- 빈 월이 있으면 "일상생활지도" 등 정기운영 프로그램 1개 이상 포함"""
    
    try:
        result = safe_gemini_json(prompt, system_instruction, max_retries=2)
        if result:
            for m in range(1, 7):
                month_key = f"{m}월"
                if month_key not in result or not result[month_key]:
                    result[month_key] = [get_default_monthly_template()]
        return result
    except ValueError:
        return None


def generate_part4(compact_text: str, month_bucket: dict = None, guideline_rules: dict = None) -> dict:
    """Part4 하반기 월별계획 (7~12월) + 예산/평가 생성. month_bucket이 있으면 해당 월에 배치."""
    
    rules = guideline_rules.get('part4', {}) if guideline_rules else {}
    monthly_rules = rules.get('monthly_program', {})
    max_programs = monthly_rules.get('max_programs_per_month', 8)
    content_max = monthly_rules.get('columns', {}).get('content', {}).get('max_chars_no_space', 200)
    budget_max = rules.get('budget_table', {}).get('max_rows', 10)
    feedback_max = rules.get('feedback_summary', {}).get('max_rows', 5)
    
    pre_bucketed = {}
    if month_bucket:
        for m in range(7, 13):
            month_key = f"{m}월"
            programs = month_bucket.get(m, [])[:max_programs]
            pre_bucketed[month_key] = [
                {
                    "program_name": p.get('program_name', ''),
                    "target": p.get('target', '전체아동'),
                    "staff": p.get('staff', '돌봄교사'),
                    "cycle": p.get('cycle', '')
                } for p in programs
            ]
    
    pre_bucket_info = ""
    if pre_bucketed:
        pre_bucket_info = "\n[월별 프로그램 배치 정보 - 반드시 이 월에 해당 프로그램을 배치하세요]:\n"
        for month, progs in pre_bucketed.items():
            if progs:
                prog_names = ", ".join([p['program_name'] for p in progs])
                pre_bucket_info += f"  {month}: {prog_names}\n"
    
    system_instruction = """당신은 사업계획 일정 및 예산 전문가입니다.
한국어, 이모지 금지, ● 기호만 사용.
출력: 오직 JSON 객체 1개만."""

    prompt = f"""다음 파일 요약을 기반으로 Part4 (하반기 7~12월 + 예산/평가) JSON을 생성하세요.

{compact_text}
{pre_bucket_info}

출력 스키마:
{{
  "monthly_plan": {{
    "7월": [{{"big_category": "보호", "mid_category": "생활", "program_name": "프로그램명", "target": "대상", "staff": "인력", "content": "● 내용"}}],
    "8월": [...],
    "9월": [...],
    "10월": [...],
    "11월": [...],
    "12월": [...]
  }},
  "budget_table": [
    {{"category": "인건비", "amount": "50,000,000원", "details": "세부내용"}}
  ],
  "feedback_summary": [
    {{"area": "보호", "problem": "문제점", "plan": "개선계획"}}
  ]
}}

[중요] 
- 각 월당 프로그램 {max_programs}개 이내
- 각 content {content_max}자 (공백 제외) 이내
- budget {budget_max}행, feedback {feedback_max}행 이내
- 빈 월이 있으면 "일상생활지도" 등 정기운영 프로그램 1개 이상 포함"""
    
    try:
        result = safe_gemini_json(prompt, system_instruction, max_retries=2)
        if result and "monthly_plan" in result:
            for m in range(7, 13):
                month_key = f"{m}월"
                if month_key not in result["monthly_plan"] or not result["monthly_plan"][month_key]:
                    result["monthly_plan"][month_key] = [get_default_monthly_template()]
        return result
    except ValueError:
        return None


def _is_bullet_format(text: str) -> bool:
    """텍스트가 불릿 형식인지 확인합니다."""
    if not text:
        return False
    lines = [l.strip() for l in text.strip().split('\n') if l.strip()]
    if not lines:
        return False
    bullet_lines = sum(1 for l in lines if l.startswith('•') or l.startswith('●'))
    return bullet_lines >= len(lines) * 0.5


def _ensure_bullet_prefix(text: str) -> str:
    """모든 줄이 '• '로 시작하도록 강제합니다."""
    if not text:
        return text
    lines = text.strip().split('\n')
    result = []
    for line in lines:
        line = line.strip()
        if not line:
            continue
        line = line.lstrip('•●-* ')
        result.append(f"• {line}")
    return '\n'.join(result)


def _ensure_bullet_count(text: str, target_count: int) -> str:
    """불릿 개수를 target_count에 맞춥니다."""
    if not text or target_count <= 0:
        return text
    
    lines = [l.strip() for l in text.strip().split('\n') if l.strip()]
    
    if len(lines) > target_count:
        lines = lines[:target_count]
    elif len(lines) < target_count:
        while len(lines) < target_count:
            lines.append("• (추가 내용 필요)")
    
    return '\n'.join(lines)


def _truncate_to_max_no_space(text: str, max_chars: int) -> str:
    """공백 제외 글자수 기준으로 자릅니다."""
    if not text or max_chars <= 0:
        return text
    
    if count_chars_no_space(text) <= max_chars:
        return text
    
    is_bullet = _is_bullet_format(text)
    
    if is_bullet:
        lines = text.strip().split('\n')
        result_lines = []
        current_count = 0
        for line in lines:
            line_count = count_chars_no_space(line)
            if current_count + line_count <= max_chars:
                result_lines.append(line)
                current_count += line_count
            else:
                remaining = max_chars - current_count
                if remaining > 10:
                    truncated = _truncate_line_to_chars(line, remaining)
                    result_lines.append(truncated)
                break
        return '\n'.join(result_lines)
    else:
        return smart_truncate(text, max_chars)


def _truncate_line_to_chars(line: str, max_chars: int) -> str:
    """한 줄을 공백 제외 max_chars 이내로 자릅니다."""
    if count_chars_no_space(line) <= max_chars:
        return line
    
    chars_counted = 0
    cut_idx = len(line)
    for i, ch in enumerate(line):
        if ch not in ' \n\t':
            chars_counted += 1
        if chars_counted >= max_chars:
            cut_idx = i + 1
            break
    
    return line[:cut_idx].rstrip() + "..."


PADDING_PHRASES = [
    "아동의 건강한 성장과 발달을 위한 체계적인 지원이 필요합니다.",
    "지역사회와 연계한 통합적 돌봄 서비스를 제공하고자 합니다.",
    "아동의 사회성 발달과 정서적 안정을 도모합니다.",
    "안전하고 건강한 돌봄 환경을 조성하여 아동의 권익을 보호합니다.",
    "맞춤형 프로그램을 통해 아동의 다양한 욕구를 충족시킵니다.",
    "전문 인력의 체계적인 관리로 양질의 서비스를 제공합니다.",
    "지속적인 모니터링과 평가를 통해 서비스 품질을 향상시킵니다.",
    "아동과 가정의 복지 증진을 위해 다각적인 지원을 실시합니다.",
]


def _pad_to_min_chars(text: str, min_chars: int, is_bullet: bool = False) -> str:
    """min_chars 미달 시 문장을 보강합니다. LLM 호출 없이 결정적으로 100% 보장."""
    if not text:
        text = "내용이 필요합니다."
    if min_chars <= 0:
        return text
    
    current = count_chars_no_space(text)
    if current >= min_chars:
        return text
    
    phrase_idx = 0
    
    if is_bullet:
        lines = [l.strip() for l in text.strip().split('\n') if l.strip()]
        if not lines:
            lines = ["• 내용이 필요합니다."]
        
        while count_chars_no_space('\n'.join(lines)) < min_chars:
            for i in range(len(lines)):
                padding = PADDING_PHRASES[phrase_idx % len(PADDING_PHRASES)]
                lines[i] = lines[i].rstrip() + " " + padding
                phrase_idx += 1
                if count_chars_no_space('\n'.join(lines)) >= min_chars:
                    break
        
        return '\n'.join(lines)
    else:
        while count_chars_no_space(text) < min_chars:
            padding = PADDING_PHRASES[phrase_idx % len(PADDING_PHRASES)]
            text = text.rstrip() + " " + padding
            phrase_idx += 1
        
        return text


def _apply_text_rule(text: str, rule: dict, field_name: str = "") -> dict:
    """텍스트 필드에 규칙을 적용합니다. 적용 결과와 로그를 반환합니다."""
    fmt = rule.get('format', 'paragraph')
    bullet_count = rule.get('bullet_count', 0)
    max_chars = rule.get('max_chars_no_space', 0)
    min_chars = rule.get('min_chars_no_space', 0)
    
    if not text:
        if fmt == 'bullet' and bullet_count > 0:
            text = '\n'.join([f"• 항목 {i+1} 내용이 필요합니다." for i in range(bullet_count)])
        else:
            text = "내용이 필요합니다."
    
    original_count = count_chars_no_space(text)
    result = text
    
    if fmt == 'bullet' and bullet_count > 0:
        result = _ensure_bullet_prefix(result)
        result = _ensure_bullet_count(result, bullet_count)
    
    if min_chars > 0:
        result = _pad_to_min_chars(result, min_chars, is_bullet=(fmt == 'bullet'))
    
    if max_chars > 0:
        result = _truncate_to_max_no_space(result, max_chars)
    
    if fmt == 'bullet' and bullet_count > 0:
        result = _ensure_bullet_prefix(result)
        result = _ensure_bullet_count(result, bullet_count)
    
    final_count = count_chars_no_space(result)
    bullet_lines = len([l for l in result.split('\n') if l.strip().startswith('•')]) if fmt == 'bullet' else 0
    
    min_ok = (min_chars == 0) or (final_count >= min_chars)
    max_ok = (max_chars == 0) or (final_count <= max_chars)
    status = "✓" if (min_ok and max_ok) else "✗"
    
    range_str = f"{min_chars}~{max_chars}" if max_chars > 0 else f"{min_chars}~∞"
    log = f"{field_name}: {original_count}→{final_count}자 ({range_str}) {status}"
    if fmt == 'bullet':
        bullet_ok = "✓" if bullet_lines == bullet_count else "✗"
        log += f", 불릿:{bullet_lines}/{bullet_count} {bullet_ok}"
    
    return {"text": result, "log": log}


def _apply_table_rule(table: list, rule: dict, table_name: str = "") -> dict:
    """테이블에 규칙을 적용합니다."""
    if not table or not isinstance(table, list):
        return {"table": [], "log": f"{table_name}: 빈 테이블"}
    
    max_rows = rule.get('max_rows', rule.get('max_rows_per_category', 100))
    columns = rule.get('columns', {})
    
    if len(table) > max_rows:
        table = table[:max_rows]
    
    for row in table:
        if isinstance(row, dict):
            for col_name, col_rule in columns.items():
                if col_name in row and row[col_name]:
                    cell_max = col_rule.get('max_chars_no_space', 0)
                    cell_min = col_rule.get('min_chars_no_space', 0)
                    cell_bullet = col_rule.get('bullet_count', 0)
                    
                    cell_text = str(row[col_name])
                    
                    if cell_bullet > 0:
                        cell_text = _ensure_bullet_prefix(cell_text)
                        cell_text = _ensure_bullet_count(cell_text, cell_bullet)
                    
                    if cell_max > 0:
                        cell_text = _truncate_to_max_no_space(cell_text, cell_max)
                    
                    if cell_min > 0:
                        cell_text = _pad_to_min_chars(cell_text, cell_min, is_bullet=(cell_bullet > 0))
                    
                    if cell_bullet > 0:
                        cell_text = _ensure_bullet_prefix(cell_text)
                    
                    row[col_name] = cell_text
    
    log = f"{table_name}: {len(table)}행 (max:{max_rows})"
    return {"table": table, "log": log}


def apply_guidelines_to_analysis(data: dict, guideline_rules: dict) -> tuple:
    """분석 결과 전체에 작성지침을 강제 적용합니다. (data, logs) 튜플 반환."""
    if not data or not guideline_rules:
        return data, ["규칙 또는 데이터 없음"]
    
    logs = []
    
    p1_rules = guideline_rules.get('part1', {})
    part1 = data.get('part1_general', {})
    
    text_fields = ['need_1_user_desire', 'need_2_1_regional', 'need_2_2_environment', 
                   'need_2_3_educational', 'purpose_text', 'goals_text']
    
    for field in text_fields:
        if field in p1_rules:
            text = part1.get(field, '')
            result = _apply_text_rule(text, p1_rules[field], field)
            part1[field] = result['text']
            logs.append(f"[Part1] {result['log']}")
    
    for table_name in ['feedback_table', 'total_review_table']:
        if table_name in p1_rules and table_name in part1:
            result = _apply_table_rule(part1[table_name], p1_rules[table_name], table_name)
            part1[table_name] = result['table']
            logs.append(f"[Part1] {result['log']}")
    
    data['part1_general'] = part1
    
    p2_rules = guideline_rules.get('part2', {})
    part2 = data.get('part2_programs', {})
    
    for category, cat_data in part2.items():
        if isinstance(cat_data, dict):
            for table_name in ['detail_table', 'eval_table']:
                if table_name in p2_rules and table_name in cat_data:
                    result = _apply_table_rule(cat_data[table_name], p2_rules[table_name], f"{category}/{table_name}")
                    cat_data[table_name] = result['table']
                    logs.append(f"[Part2] {result['log']}")
    
    data['part2_programs'] = part2
    
    p3_rules = guideline_rules.get('part3', {})
    part3 = data.get('part3_monthly_plan', {})
    monthly_rule = p3_rules.get('monthly_program', {})
    max_per_month = monthly_rule.get('max_programs_per_month', 8)
    columns = monthly_rule.get('columns', {})
    
    for month_key, programs in part3.items():
        if isinstance(programs, list):
            if len(programs) > max_per_month:
                programs = programs[:max_per_month]
            
            for prog in programs:
                if isinstance(prog, dict):
                    for col_name, col_rule in columns.items():
                        if col_name in prog and prog[col_name]:
                            cell_max = col_rule.get('max_chars_no_space', 0)
                            cell_min = col_rule.get('min_chars_no_space', 0)
                            cell_bullet = col_rule.get('bullet_count', 0)
                            cell_text = str(prog[col_name])
                            
                            if cell_bullet > 0:
                                cell_text = _ensure_bullet_prefix(cell_text)
                                cell_text = _ensure_bullet_count(cell_text, cell_bullet)
                            
                            if cell_max > 0:
                                cell_text = _truncate_to_max_no_space(cell_text, cell_max)
                            
                            if cell_min > 0:
                                cell_text = _pad_to_min_chars(cell_text, cell_min, is_bullet=(cell_bullet > 0))
                            
                            if cell_bullet > 0:
                                cell_text = _ensure_bullet_prefix(cell_text)
                            
                            prog[col_name] = cell_text
            
            part3[month_key] = programs
            logs.append(f"[Part3] {month_key}: {len(programs)}개 (max:{max_per_month})")
    
    data['part3_monthly_plan'] = part3
    
    p4_rules = guideline_rules.get('part4', {})
    part4_monthly = data.get('part4_monthly_plan', {})
    monthly_rule_p4 = p4_rules.get('monthly_program', {})
    max_per_month_p4 = monthly_rule_p4.get('max_programs_per_month', 8)
    columns_p4 = monthly_rule_p4.get('columns', {})
    
    for month_key, programs in part4_monthly.items():
        if isinstance(programs, list):
            if len(programs) > max_per_month_p4:
                programs = programs[:max_per_month_p4]
            
            for prog in programs:
                if isinstance(prog, dict):
                    for col_name, col_rule in columns_p4.items():
                        if col_name in prog and prog[col_name]:
                            cell_max = col_rule.get('max_chars_no_space', 0)
                            cell_min = col_rule.get('min_chars_no_space', 0)
                            cell_bullet = col_rule.get('bullet_count', 0)
                            cell_text = str(prog[col_name])
                            
                            if cell_bullet > 0:
                                cell_text = _ensure_bullet_prefix(cell_text)
                                cell_text = _ensure_bullet_count(cell_text, cell_bullet)
                            
                            if cell_max > 0:
                                cell_text = _truncate_to_max_no_space(cell_text, cell_max)
                            
                            if cell_min > 0:
                                cell_text = _pad_to_min_chars(cell_text, cell_min, is_bullet=(cell_bullet > 0))
                            
                            if cell_bullet > 0:
                                cell_text = _ensure_bullet_prefix(cell_text)
                            
                            prog[col_name] = cell_text
            
            part4_monthly[month_key] = programs
            logs.append(f"[Part4] {month_key}: {len(programs)}개 (max:{max_per_month_p4})")
    
    data['part4_monthly_plan'] = part4_monthly
    
    budget_eval = data.get('part4_budget_evaluation', {})
    
    if 'budget_table' in p4_rules and 'budget_table' in budget_eval:
        result = _apply_table_rule(budget_eval['budget_table'], p4_rules['budget_table'], 'budget_table')
        budget_eval['budget_table'] = result['table']
        logs.append(f"[Part4] {result['log']}")
    
    if 'feedback_summary' in p4_rules and 'feedback_summary' in budget_eval:
        result = _apply_table_rule(budget_eval['feedback_summary'], p4_rules['feedback_summary'], 'feedback_summary')
        budget_eval['feedback_summary'] = result['table']
        logs.append(f"[Part4] {result['log']}")
    
    data['part4_budget_evaluation'] = budget_eval
    
    return data, logs


def get_partitioned_analysis(compact_text: str, progress_callback=None, month_bucket: dict = None, guideline_rules: dict = None) -> dict:
    """파트별로 나눠서 Gemini 분석 수행. 부분 실패 허용. month_bucket과 guideline_rules 적용."""
    result = {
        "part1_general": {},
        "part2_programs": {},
        "part3_monthly_plan": {},
        "part4_monthly_plan": {},
        "part4_budget_evaluation": {"budget_table": [], "feedback_summary": []},
        "_failed_parts": [],
        "_guideline_logs": []
    }
    
    if progress_callback:
        progress_callback("Part 1 (총괄/기획) 생성 중...")
    part1 = generate_part1(compact_text)
    if part1:
        result["part1_general"] = part1
    else:
        result["_failed_parts"].append("part1")
    
    if progress_callback:
        progress_callback("Part 2 (세부사업) 생성 중...")
    part2 = generate_part2(compact_text)
    if part2:
        result["part2_programs"] = part2
    else:
        result["_failed_parts"].append("part2")
    
    if progress_callback:
        progress_callback("Part 3 (상반기 월별계획) 생성 중...")
    part3 = generate_part3(compact_text, month_bucket=month_bucket, guideline_rules=guideline_rules)
    if part3:
        result["part3_monthly_plan"] = part3
    else:
        result["_failed_parts"].append("part3")
    
    if progress_callback:
        progress_callback("Part 4 (하반기 + 예산/평가) 생성 중...")
    part4 = generate_part4(compact_text, month_bucket=month_bucket, guideline_rules=guideline_rules)
    if part4:
        if "monthly_plan" in part4:
            result["part4_monthly_plan"] = part4["monthly_plan"]
        if "budget_table" in part4:
            result["part4_budget_evaluation"]["budget_table"] = part4["budget_table"]
        if "feedback_summary" in part4:
            result["part4_budget_evaluation"]["feedback_summary"] = part4["feedback_summary"]
    else:
        result["_failed_parts"].append("part4")
    
    if guideline_rules:
        if progress_callback:
            progress_callback("작성지침 강제 적용 중...")
        result, guideline_logs = apply_guidelines_to_analysis(result, guideline_rules)
        result["_guideline_logs"] = guideline_logs
    
    return result


def parse_json_response(response_text: str) -> dict:
    """Parse JSON from Gemini response, removing markdown code blocks if present."""
    return _extract_json_from_text(response_text)


def get_gemini_analysis(text: str) -> dict:
    """Analyze text using Gemini and return structured JSON data."""
    if not GEMINI_API_KEY:
        st.error("GEMINI_API_KEY가 설정되지 않았습니다. Replit Secrets에 GEMINI_API_KEY를 추가하세요.")
        return None
    
    print(f"[Gemini] model={GEMINI_MODEL}")
    
    system_instruction = """당신은 연간 사업 평가 문서를 분석하는 전문가입니다.

**절대 규칙 (반드시 준수)**: 
- 출력은 반드시 JSON 객체({}) 1개로만 한다
- 절대 배열([])로 시작하지 않는다
- 설명문, 마크다운, 코드블록(```) 금지
- 첫 문자는 반드시 '{' 이어야 한다
- JSON 외 어떤 텍스트도 포함하지 않는다

주어진 문서를 분석하여 반드시 아래의 정확한 JSON 구조로 응답해주세요.

**중요: 모든 내용, 요약, 테이블 값은 반드시 한국어(한글)로 작성해야 합니다.**
**문체: 전문적이고 공식적인 행정 문서 스타일의 한국어를 사용하세요.**

필수 JSON 구조:
{
  "part1_general": {
    "need_1_user_desire": "이용아동의 욕구 및 문제점 분석 내용 (700~800자로 상세히 작성, 공백 제외 500자 이상 필수)",
    "need_2_1_regional": "지역적 특성 - 해당 지역의 인구통계, 경제적 특성, 행정구역 특성 등 (700~800자로 상세히 작성)",
    "need_2_2_environment": "주변환경 - 시설 주변의 교통, 접근성, 주거환경, 편의시설 등 (700~800자로 상세히 작성)",
    "need_2_3_educational": "교육적 특성 - 지역 내 학교, 교육기관, 학습환경 특성 등 (700~800자로 상세히 작성)",
    "feedback_table": [
      {"area": "보호", "problem": "• [문제점1: 약 80자 상세 서술]\\n\\n• [문제점2: 약 80자 상세 서술]\\n\\n• [문제점3: 약 80자 상세 서술]", "improvement": "• [개선방안1: 최소 100자 이상, 누가(Who), 언제(When), 무엇을(What), 어떻게(How) 포함하여 구체적으로 작성]\\n\\n• [개선방안2: 최소 100자 이상, 누가(Who), 언제(When), 무엇을(What), 어떻게(How) 포함하여 구체적으로 작성]\\n\\n• [개선방안3: 최소 100자 이상, 누가(Who), 언제(When), 무엇을(What), 어떻게(How) 포함하여 구체적으로 작성]"},
      {"area": "교육", "problem": "...", "improvement": "..."},
      {"area": "문화", "problem": "...", "improvement": "..."},
      {"area": "정서지원", "problem": "...", "improvement": "..."},
      {"area": "지역사회연계", "problem": "...", "improvement": "..."}
    ],
    "total_review_table": [
      {"category": "운영평가", "content": "● **시설운영 현황**: 상세설명(300자 이상)...\\n\\n● **인력관리 성과**: 상세설명(300자 이상)...\\n\\n● **예산집행 분석**: 상세설명(300자 이상)... (각 항목당 최소 300자 이상, 총 약 1000자)"},
      {"category": "아동평가", "content": "● **이용현황 분석**: 상세설명(300자 이상)...\\n\\n● **발달성과 평가**: 상세설명(300자 이상)...\\n\\n● **만족도 결과**: 상세설명(300자 이상)... (각 항목당 최소 300자 이상, 총 약 1000자)"},
      {"category": "프로그램평가", "content": "● **프로그램 운영실적**: 상세설명(300자 이상)...\\n\\n● **목표달성 분석**: 상세설명(300자 이상)...\\n\\n● **개선필요 사항**: 상세설명(300자 이상)... (각 항목당 최소 300자 이상, 총 약 1000자)"},
      {"category": "후원활동측면", "content": "● **후원현황 분석**: 상세설명(300자 이상)...\\n\\n● **기업연계 성과**: 상세설명(300자 이상)...\\n\\n● **네트워크 구축**: 상세설명(300자 이상)... (각 항목당 최소 300자 이상, 총 약 1000자)"},
      {"category": "환류방안", "content": "● **차년도 목표설정**: 상세설명(300자 이상)...\\n\\n● **개선계획 수립**: 상세설명(300자 이상)...\\n\\n● **실행방안 도출**: 상세설명(300자 이상)... (각 항목당 최소 300자 이상, 총 약 1000자)"}
    ],
    "satisfaction_survey": {
      "total_respondents": 30,
      "survey_data": [
        {"문항": "1. 급식 및 간식의 질과 위생 상태", "5점": 15, "4점": 10, "3점": 3, "2점": 1, "1점": 1},
        {"문항": "2. 프로그램 내용과 운영", "5점": 14, "4점": 11, "3점": 3, "2점": 1, "1점": 1},
        {"문항": "3. 시설의 안전 및 위생 관리", "5점": 16, "4점": 9, "3점": 3, "2점": 1, "1점": 1},
        {"문항": "4. 담당 선생님의 아동 지도 능력", "5점": 15, "4점": 10, "3점": 3, "2점": 1, "1점": 1},
        {"문항": "5. 학습지도 프로그램의 효과", "5점": 12, "4점": 12, "3점": 4, "2점": 1, "1점": 1},
        {"문항": "6. 아동의 정서적 지원", "5점": 13, "4점": 11, "3점": 4, "2점": 1, "1점": 1},
        {"문항": "7. 부모 상담 및 소통", "5점": 11, "4점": 12, "3점": 5, "2점": 1, "1점": 1},
        {"문항": "8. 문화체험 활동", "5점": 10, "4점": 11, "3점": 6, "2점": 2, "1점": 1},
        {"문항": "9. 시설 운영 시간", "5점": 11, "4점": 11, "3점": 5, "2점": 2, "1점": 1},
        {"문항": "10. 전반적인 시설 운영", "5점": 14, "4점": 11, "3점": 3, "2점": 1, "1점": 1}
      ],
      "subjective_question": "기타 건의사항 및 개선 의견",
      "subjective_analysis": "주관식 응답 분석 내용 (500자 이상으로 상세히 작성)",
      "overall_suggestion": "만족도 조사 결과에 대한 종합 분석 및 제언 (500자 이상으로 상세히 작성)"
    },
    "purpose_text": "지역사회 내 돌봄이 필요한 만 18세 미만 아동을 대상으로 가정의 기능을 보완하고 사회적 안전망을 강화하여 차별 없는 성장 환경을 조성함을 목적으로 하는 사업목적 (약 500자의 단일 응집 문단으로 작성)",
    "goals_text": "● **건강하고 안전한 권리 기반 돌봄 환경 조성**: (300자 이상) 일상생활 및 위생 관리...\\n\\n● **전인적 성장을 위한 맞춤형 교육 지원**: (300자 이상) 기초 학습 지도...\\n\\n● **다양한 문화체험을 통한 사회성 함양**: (300자 이상) 캠프 및 체험활동...\\n\\n● **정서적 안정과 가정기능 회복 지원**: (300자 이상) 상담 및 사례관리...\\n\\n● **지속 가능한 운영과 지역사회 연계 강화**: (300자 이상) 자원봉사 및 후원관리..."
  },
  "part2_programs": {
    "보호": {
      "subcategories": ["생활", "안전", "가족기능강화"],
      "detail_table": [
        {"sub_area": "생활/안전/가족기능강화 중 택일", "program_name": "프로그램명", "expected_effect": "기대효과 (100자 이상)", "target": "대상아동", "count": "계획인원", "cycle": "주기", "content": "● **세부내용1**: 상세 설명...\\n● **세부내용2**: 상세 설명..."}
      ],
      "eval_table": [
        {"sub_area": "세부영역명", "program_name": "프로그램명", "expected_effect": "● **효과1**: 상세 설명(아동의 행동, 태도, 기술 변화를 구체적으로 기술)...\\n● **효과2**: 상세 설명...", "main_plan": "주요 실행 계획 요약 (어떻게 진행할 것인지)", "eval_method": "프로그램 일지, 만족도 조사, 출석부, 관찰일지, 사전사후 검사 중 선택"}
      ]
    },
    "교육": {
      "subcategories": ["성장과권리", "학습", "특기적성"],
      "detail_table": [],
      "eval_table": []
    },
    "문화": {
      "subcategories": ["체험활동"],
      "detail_table": [],
      "eval_table": []
    },
    "정서지원": {
      "subcategories": ["상담"],
      "detail_table": [],
      "eval_table": []
    },
    "지역사회연계": {
      "subcategories": ["연계"],
      "detail_table": [],
      "eval_table": []
    }
  },
  "part3_monthly_plan": {
    "1월": [
      {"big_category": "대분류", "mid_category": "중분류", "program_name": "프로그램명", "target": "참여자", "staff": "수행인력", "content": "● **세부내용**: 설명..."}
    ],
    "2월": [],
    "3월": [],
    "4월": [],
    "5월": [],
    "6월": [],
    "7월": [],
    "8월": [],
    "9월": [],
    "10월": [],
    "11월": [],
    "12월": []
  },
  "part4_budget_evaluation": {
    "budget_table": [
      {"category": "인건비", "amount": "50,000,000원", "details": "돌봄교사 4명, 학습지도사 2명, 상담사 1명 인건비"}
    ],
    "feedback_summary": [
      {"area": "보호", "problem": "문제점 요약", "plan": "개선계획 요약"}
    ]
  }
}

중요 사항:
- **모든 텍스트 값은 반드시 한국어(한글)로 작성하세요.**
- **문체는 전문적이고 공식적인 행정 문서 스타일을 사용하세요.**
- JSON 키(key)는 영어로 유지하되, 값(value)은 모두 한국어로 작성하세요.
- **satisfaction_survey 작성 규칙:**
  - total_respondents: 총 응답자 수 (정수, 기본값 30명 또는 문서에서 추출)
  - survey_data: 10개 문항에 대해 5점~1점 척도별 응답 인원수를 배분하세요.
    - 형식: {"문항": "1. 질문내용", "5점": 숫자, "4점": 숫자, "3점": 숫자, "2점": 숫자, "1점": 숫자}
    - 각 문항의 응답 인원 합계가 total_respondents와 일치해야 합니다.
    - 긍정적인 결과를 위해 5점과 4점에 응답이 많이 분포되도록 현실적으로 배분하세요.
  - 문항은 업로드된 문서에서 추출하거나, 없으면 아동복지시설 평가 표준 항목(급식, 프로그램, 안전, 교사, 학습지도, 정서지원, 소통, 문화활동, 운영시간, 전반적 만족도) 10개를 사용하세요.
  - subjective_analysis: 주관식 응답의 주요 내용과 경향을 500자 이상으로 상세히 분석하세요.
  - overall_suggestion: 만족도 결과를 바탕으로 종합 분석 및 개선 제언을 500자 이상으로 상세히 작성하세요.
- 문서에 정보가 없는 경우 적절한 한국어 기본값을 사용하세요.
- **need_1_user_desire는 700~800자 분량(공백 제외 500자 이상)으로 이용아동의 욕구와 문제점을 상세히 분석하세요. 간략히 요약하지 말고 구체적인 예시와 근거를 들어 충분히 서술하세요.**
- **need_2_1_regional, need_2_2_environment, need_2_3_educational 각각 700~800자 분량으로 상세히 작성하세요. 간략히 요약하지 말고 구체적인 예시와 근거를 들어 충분히 서술하세요.**
- **feedback_table은 반드시 5개 영역(보호, 교육, 문화, 정서지원, 지역사회연계) 순서로 작성하세요.**
  - 각 영역당 문제점과 개선방안을 각각 3개씩 bullet point(•)로 구분하여 작성하세요.
  - 각 bullet point는 최소 150자 이상으로 상세히 서술하세요.
  - 문제점에는 구체적인 원인(Why)을, 개선방안에는 구체적인 방법(How)을 포함하세요.
  - 문제점의 1번 항목과 개선방안의 1번 항목이 서로 대응되어야 합니다 (1:1 매핑).
  - **[문제점 필수 규칙 - 형식]**:
    - 문제점(problem)의 각 bullet point 사이에 **반드시 빈 줄(\\n\\n)을 삽입**하여 개선방안과 동일한 형식으로 시각적 정렬을 맞추세요.
    - 형식 예시:
      "• [문제점1 내용...]\\n\\n• [문제점2 내용...]\\n\\n• [문제점3 내용...]"
  - **[개선방안 필수 규칙 - 길이와 형식]**:
    - 개선방안(improvement)의 각 bullet point는 반드시 **최소 100자 이상**(한글 기준)으로 작성하세요. 문제점보다 반드시 더 길게 작성해야 합니다.
    - 각 개선방안에는 **누가(Who), 언제(When), 무엇을(What), 어떻게(How)**를 구체적으로 명시하세요.
      예시: "2025년 상반기까지 돌봄 담당 인력 2명을 추가 채용하여 아동 1인당 돌봄 비율을 개선하고..."
    - 각 bullet point 사이에 **반드시 빈 줄(\\n\\n)을 삽입**하여 시각적으로 분리하세요.
    - 형식 예시:
      "• [개선방안1 내용 100자 이상...]\\n\\n• [개선방안2 내용 100자 이상...]\\n\\n• [개선방안3 내용 100자 이상...]"
- **total_review_table은 반드시 5개 영역(운영평가, 아동평가, 프로그램평가, 후원활동측면, 환류방안) 순서로 작성하세요.**
  - 각 영역당 최소 1000자 이상으로 매우 상세하고 전문적으로 서술하세요.
  - 요약하지 말고, 구체적인 수치, 사례, 성과, 과제를 빠짐없이 포함하여 작성하세요.
  - **[가독성 향상 필수 규칙 - 블랙 서클 서식]**:
    - 절대로 줄글(wall of text)로 작성하지 마세요!
    - **이모지 사용 금지!** 반드시 ● (검정 원) 기호만 사용하세요.
    - 반드시 ● **[핵심키워드]**: [상세설명] 형식으로 구조화하세요.
    - **프로그램평가, 환류방안 영역은 각 항목당 최소 300자 이상**으로 상세히 작성하세요.
    - 각 소제목 사이에 빈 줄(\\n\\n)을 삽입하여 시각적으로 분리하세요.
    - 예시 형식:
      "● **시설운영 현황**: 2024년 시설 운영률 95%를 달성하였으며, 일일 평균 이용 아동 수는 45명으로 정원 대비 90%의 높은 이용률을 기록하였습니다. 월별 운영 현황을 분석한 결과 방학 기간 이용률이 평시 대비 15% 증가하는 경향을 보였으며, 이에 따라 방학 특별 프로그램을 강화하여 운영하였습니다. 시설 운영 시간은 평일 13시부터 19시까지 6시간 운영하였으며, 연장 돌봄 요청에 따라 20시까지 운영 시간을 확대하는 방안을 검토 중입니다...\\n\\n● **인력관리 성과**: 전문인력 8명을 안정적으로 운영하며..."
    - 각 영역당 최소 3개 이상의 소제목으로 구조화하세요.
- **need_1_user_desire, need_2_1_regional, need_2_2_environment, need_2_3_educational, subjective_analysis, overall_suggestion 작성 시에도 동일하게 블랙 서클 서식을 적용하세요.**
  - **이모지 사용 금지!** 반드시 ● (검정 원) 기호만 사용하세요.
  - 형식 예시: "● **현황분석**: 설명내용...\\n\\n● **핵심과제**: 설명내용...\\n\\n● **시사점**: 설명내용..."
- **사업목적(purpose_text) 작성 규칙:**
  - 약 500자의 단일 응집 문단으로 작성하세요.
  - 형식적이고 제안서 스타일의 언어를 사용하세요 (예: "본 사업은 ~함을 목적으로 한다.").
  - 다음 핵심 개념을 통합하여 작성하세요: 만 18세 미만 취약계층 아동, 가정 기능 보완, 사회적 안전망 강화, 차별 없는 성장 환경 조성.
  - 이모지 사용 금지. 공식적인 행정 문서 스타일로 작성하세요.
- **사업목표(goals_text) 작성 규칙:**
  - 반드시 5개의 별도 bullet point로 구성하세요.
  - 형식: ● **[핵심키워드/주제]**: [상세 내용...]
  - **각 bullet point는 최소 300자 이상**으로 상세히 작성하세요.
  - 5개 목표 구성:
    1. **건강하고 안전한 권리 기반 돌봄 환경 조성**: 일상생활 및 위생관리, 균형 잡힌 급식, 아동 권리 및 인권 교육, 아동 참여권 보장 내용 포함
    2. **전인적 성장을 위한 맞춤형 교육 지원**: 기초 학습 지도, 문해력 향상 독서 지도, 예술/체육 교육, 긍정적 강화를 통한 교육 격차 해소 내용 포함
    3. **다양한 문화체험을 통한 사회성 함양**: 캠프 및 체험활동, 사회성 발달, 정서적 해소 내용 포함
    4. **정서적 안정과 가정기능 회복 지원**: 상담, 사례관리, 가족 지원, 심리적 안정 및 가정기능 회복 내용 포함
    5. **지속 가능한 운영과 지역사회 연계 강화**: 자원봉사자/후원자 관리, 투명한 운영, 지역사회 네트워크 구축 내용 포함
  - 각 bullet point 사이에 빈 줄(\\n\\n)을 삽입하세요.
  - 이모지 사용 금지. ● 기호만 사용하세요.
- **part2_programs 작성 규칙 (고정 세부영역 분류 필수):**
  - 5개 주요 카테고리(보호, 교육, 문화, 정서지원, 지역사회연계)별로 프로그램을 분류하세요.
  - 각 카테고리에는 고정된 세부영역(sub_area)이 있으며, 프로그램을 반드시 해당 세부영역에 배치하세요:
    ● **보호**: 생활(위생/급식/일상지도), 안전(야간보호/귀가지도/안전교육), 가족기능강화(부모교육/가족캠프)
    ● **교육**: 성장과권리(아동권리교육/자치회의), 학습(기초학력지도/독서/영어), 특기적성(음악/미술/스포츠)
    ● **문화**: 체험활동(캠프/현장학습/공연관람)
    ● **정서지원**: 상담(개별상담/집단상담/보호자상담)
    ● **지역사회연계**: 연계(자원봉사관리/후원연계/기관협력)
  - detail_table 필수 필드: sub_area, program_name, expected_effect, target, count, cycle, content
  - **expected_effect**: 100자 이상의 상세한 기대효과 작성
  - **content (계획내용)**: ● 불릿 포인트 사용, 각 내용은 상세히 작성 (예: "● **급식 관리**: 균형 잡힌 영양 식단을 주간 단위로 계획하여...")
  - 문서에서 프로그램이 추출되지 않은 세부영역의 경우, 지역아동센터 표준 프로그램을 현실적으로 생성하세요.
  - 이모지 사용 금지, ● 기호만 사용하세요.
- **eval_table (평가계획) 작성 규칙:**
  - eval_table은 detail_table에서 정의된 프로그램을 그대로 연계하여 작성하세요. 새로운 프로그램을 만들지 마세요.
  - 각 카테고리(보호, 교육, 문화, 정서지원, 지역사회연계)별로 해당 프로그램에 대한 평가계획을 작성하세요.
  - eval_table 필수 필드: sub_area, program_name, expected_effect, main_plan, eval_method
  - **sub_area (세부영역)**: detail_table의 세부영역과 동일하게 작성
  - **program_name (프로그램명)**: detail_table의 프로그램명과 동일하게 작성
  - **expected_effect (기대효과) - 핵심 작성 규칙:**
    - detail_table의 기대효과보다 **더 상세하고 구체적으로** 작성해야 합니다.
    - 아동의 **행동, 태도, 기술의 구체적인 변화**를 중점적으로 기술하세요.
    - ● 불릿 포인트를 사용하여 2-3개의 구분된 효과를 작성하세요.
    - 나쁜 예: "안전성 향상"
    - 좋은 예: "● **위기대응역량 강화**: 참여 아동 100%가 모의 훈련에서 CPR 기법을 정확히 수행할 수 있게 됨\\n● **안전습관 정착**: 반복 훈련을 통해 일상생활에서의 안전 습관 준수율이 20% 향상됨"
    - 각 효과에는 **굵은 키워드**를 포함하고, 구체적인 성과 지표(예: 퍼센트, 인원, 횟수)를 명시하세요.
  - **main_plan (주요계획)**: 프로그램을 어떻게 실행할 것인지 간략히 요약 (예: "월 1회 소방서 연계 화재대피 훈련 실시")
  - **eval_method (평가방법)**: 다음 평가 도구 중 적합한 것을 선택하여 작성:
    - 프로그램 일지
    - 만족도 조사
    - 출석부
    - 관찰일지
    - 사전사후 검사
    - 여러 도구를 조합할 수 있음 (예: "프로그램 일지, 만족도 조사")
- **월별 계획(part3_monthly_plan) 작성 규칙 (지능형 날짜 배분):**
  - **핵심 로직**: part2_programs에서 추출된 모든 프로그램을 분석하고, 각 프로그램의 '주기(cycle)'를 기반으로 해당 월에 자동 배분합니다.
  - **날짜 확장 규칙 (CRITICAL)**:
    - "연중" 또는 "매월" 또는 "주 N회" → 1월~12월 전체에 배분
    - "3월~11월" → 3월, 4월, 5월, 6월, 7월, 8월, 9월, 10월, 11월에 배분
    - "여름방학" → 7월, 8월에 배분
    - "겨울방학" → 1월, 2월, 12월에 배분
    - "분기 1회" → 3월, 6월, 9월, 12월에 배분
    - "연 2회" → 1월, 7월에 배분 (상반기/하반기 대표)
    - 특정 날짜 "4월 15일" → 4월에만 배분
  - **JSON 구조**: 키는 "1월"..."12월"이고, 각 월에는 프로그램 목록이 배열로 포함됩니다.
  - **프로그램 항목 필드**: 각 프로그램은 다음 6개 필드를 포함:
    - big_category: 대분류 (보호/교육/문화/정서지원/지역사회연계)
    - mid_category: 중분류 (세부영역)
    - program_name: 프로그램명
    - target: 참여자/대상
    - staff: 수행인력 (알 수 없으면 "사회복지사", "센터장", "돌봄교사" 중 추론)
    - content: 사업내용 (● 불릿 포인트 사용, 예: "● **안전관리**: 화재대피 훈련 실시...")
  - **이모지 사용 금지**: ● 기호만 사용하세요.
- **예산 및 평가(part4_budget_evaluation) 작성 규칙:**
  - **budget_table (예산계획)**: 현실적인 예산 테이블을 생성합니다.
    - 필드: category(항목), amount(금액), details(세부내용)
    - 주요 항목: 인건비, 사업비, 운영비, 기타경비 등
    - 총계가 합리적으로 맞아야 함
  - **feedback_summary (평가 및 환류)**: 차년도 계획 환류 요약
    - 필드: area(영역), problem(문제점), plan(개선계획)
    - 5대 영역별(보호, 교육, 문화, 정서지원, 지역사회연계) 환류 요약
"""

    prompt = f"다음 문서를 분석하고 지정된 JSON 형식으로 결과를 반환해주세요:\n\n{text}"
    
    try:
        parsed = safe_gemini_json(prompt, system_instruction, max_retries=2)
        
        parsed.setdefault("part1_general", {})
        parsed.setdefault("part2_programs", {})
        parsed.setdefault("part3_monthly_plan", {})
        parsed.setdefault("part4_monthly_plan", {})
        parsed.setdefault("part4_budget_evaluation", {"budget_table": [], "feedback_summary": []})
        
        if "part1_general" in parsed and isinstance(parsed["part1_general"], dict):
            p1 = parsed["part1_general"]
            if "feedback_table" in p1:
                p1["feedback_table"] = normalize_table_rows(
                    p1["feedback_table"], 
                    ["area", "problem", "improvement"]
                )
            if "total_review_table" in p1:
                p1["total_review_table"] = normalize_table_rows(
                    p1["total_review_table"], 
                    ["category", "content"]
                )
        
        return parsed
            
    except ValueError as e:
        error_msg = str(e)
        st.error("JSON 파싱에 실패했습니다.")
        if "원문:" in error_msg:
            raw_preview = error_msg.split("원문:")[-1][:1500]
            st.code(raw_preview, language="text")
        return None
    except Exception as e:
        error_msg = str(e)
        if "404" in error_msg:
            st.error(f"Gemini API 오류: 모델 '{GEMINI_MODEL}'을 찾을 수 없습니다.")
        elif "403" in error_msg or "permission" in error_msg.lower():
            st.error("Gemini API 오류: API 키 권한을 확인하세요.")
        elif "429" in error_msg or "quota" in error_msg.lower():
            st.error("Gemini API 오류: API 쿼터를 초과했습니다.")
        else:
            st.error(f"Gemini API 오류: {error_msg}")
        return None


def get_default_data() -> dict:
    """Return default data structure for testing without API."""
    return {
        "part1_general": {
            "need_1_user_desire": "● **욕구조사 현황**: 이용아동들은 안정적인 돌봄 환경과 다양한 교육 기회를 원하고 있습니다. 특히 방과후 시간대의 체계적인 프로그램과 정서적 지원에 대한 욕구가 높게 나타났습니다.\n\n● **주요 욕구**: 학업 지원과 함께 또래 관계 형성, 자기표현 능력 향상에 대한 요구도 확인되었습니다. 최근 실시한 이용아동 대상 욕구조사 결과, 전체 응답자의 78%가 다양한 체험활동 프로그램 확대를 희망하였으며, 65%의 아동이 개별 맞춤형 학습지원 서비스를 필요로 하는 것으로 나타났습니다.\n\n● **시사점**: 정서적 어려움을 호소하는 아동 비율이 전년 대비 12% 증가하여 전문적인 심리상담 서비스에 대한 수요가 증대되고 있습니다. 코로나19 이후 사회성 발달 지연 문제가 대두되면서 또래관계 형성 프로그램에 대한 요구도 높아지고 있는 실정입니다.",
            "need_2_1_regional": "● **인구현황**: 본 시설이 위치한 지역은 수도권 외곽의 신도시로서 최근 10년간 급격한 인구 증가를 경험하였습니다. 총 인구 약 35만 명 중 0-18세 아동·청소년 인구가 약 7만 명으로 전체의 20%를 차지하고 있습니다.\n\n● **가정형태**: 지역 내 맞벌이 가정 비율이 68%에 달하며, 한부모 가정도 전국 평균 대비 1.5배 높은 수준입니다. 경제활동인구의 대부분이 서울 및 인근 대도시로 출퇴근하는 베드타운 성격이 강하여, 아동 돌봄 공백 문제가 심각한 지역적 특성을 보이고 있습니다.\n\n● **경제현황**: 지역 내 제조업 및 서비스업 종사자 비율이 높으며, 평균 가구소득은 전국 중위소득의 95% 수준입니다.",
            "need_2_2_environment": "● **교육시설**: 시설 반경 1km 이내에 초등학교 3개교, 중학교 2개교가 위치하여 이용 아동의 접근성이 매우 양호합니다.\n\n● **대중교통**: 대중교통 인프라가 잘 구축되어 있어 버스 정류장이 도보 5분 거리에 있으며, 지하철역도 10분 거리에 위치하고 있습니다. 주변에 대형마트, 병원, 공원 등 생활편의시설이 충분히 갖추어져 있어 다양한 체험활동 연계가 용이합니다.\n\n● **안전사항**: 시설 주변 도로의 교통량이 많아 아동 안전에 대한 각별한 주의가 필요하며, 통학 시간대 교통안전지도 강화가 요구됩니다. 주거 형태는 아파트 단지가 85%를 차지하며, 소규모 주택 밀집 지역도 일부 존재합니다.",
            "need_2_3_educational": "● **교육환경**: 지역 내 교육열이 높아 사교육 참여율이 전국 평균 대비 15% 높은 수준입니다. 그러나 저소득층 가정의 경우 사교육비 부담으로 인한 교육 격차 문제가 발생하고 있습니다.\n\n● **교육자원**: 지역 교육지원청에서 운영하는 방과후학교 프로그램이 다양하게 운영되고 있으나, 수요 대비 공급이 부족한 실정입니다. 인근에 도서관 2개소, 청소년수련관 1개소가 위치하여 교육 연계 자원이 풍부합니다.\n\n● **다문화지원**: 최근 지역 내 다문화 가정이 증가하면서 한국어 교육 및 문화 적응 프로그램에 대한 수요도 늘어나고 있습니다. 또한 특수교육 대상 아동을 위한 전문 교육 서비스가 부족하여 통합교육 환경 조성이 필요한 상황입니다.",
            "feedback_table": [
                {"area": "보호", "problem": "• 저녁 돌봄 시간이 18시까지로 제한되어 맞벌이 가정의 퇴근 시간과 맞지 않아 돌봄 공백이 발생하고 있음\n\n• 돌봄 공간의 안전시설(CCTV, 안전문 등)이 노후화되어 정기적인 점검과 교체가 필요한 상황임\n\n• 아동 1인당 돌봄 인력 비율이 높아 개별 아동에 대한 세심한 관찰과 지도가 어려운 실정임", "improvement": "• 2025년 상반기까지 시설 운영팀에서 저녁 돌봄 시간을 기존 18시에서 20시까지 연장 운영하고, 야간 돌봄 전담 인력 2명을 추가 채용하여 맞벌이 가정의 돌봄 공백 문제를 해결하며, 연장 돌봄 신청 절차를 간소화하여 이용 편의성을 높임\n\n• 2025년 1분기 중 시설관리 담당자가 노후화된 CCTV 8대와 안전문 3개소를 전면 교체하고, 매월 1회 정기 안전점검을 실시하며, 점검 결과를 문서화하여 관리함으로써 시설 안전성을 체계적으로 확보함\n\n• 2025년 상반기까지 돌봄 담당 인력 2명을 추가 채용하여 아동 대 돌봄 인력 비율을 10:1에서 7:1로 개선하고, 소그룹 돌봄 체계(5~7명 단위)를 도입하여 개별 아동에 대한 세심한 관찰과 맞춤형 지도를 실현함"},
                {"area": "교육", "problem": "• 온라인 학습을 위한 디지털 기기(태블릿, 노트북)가 부족하여 아동들의 디지털 리터러시 교육에 제한이 있음\n\n• 학습 지도 프로그램이 주로 국어, 수학에 집중되어 있어 영어, 과학 등 다양한 교과 지원이 부족한 상황임\n\n• 학습 부진 아동을 위한 개별 맞춤형 지도 체계가 미흡하여 학습 격차 해소에 어려움이 있음", "improvement": "• 2025년 1분기까지 교육 담당팀에서 태블릿 PC 20대를 추가 구입하고, AI 기반 스마트 학습 플랫폼을 도입하여 아동 1인당 1기기 사용이 가능한 디지털 학습 환경을 구축하며, 디지털 리터러시 교육을 주 2회 정기 운영함\n\n• 2025년 상반기까지 영어 전문 학습지도사 1명과 과학 전문 학습지도사 1명을 채용하고, 교과별 특성화 프로그램(영어 회화반, 과학 실험반)을 개발하여 주 3회 운영함으로써 교육 영역의 다양성을 확보함\n\n• 2025년 1분기부터 교육 담당자가 표준화된 학습 진단 도구를 도입하여 분기별 학습 수준을 평가하고, 학습 부진 아동 대상 개별 학습 계획(ILP)을 수립하여 주 2회 1:1 맞춤형 학습 지원을 제공함"},
                {"area": "문화", "problem": "• 문화체험 프로그램이 연간 4회 수준으로 다른 시설 대비 부족하고 프로그램 종류도 제한적인 상황임\n\n• 예술(음악, 미술) 관련 전문 강사가 부재하여 양질의 예술 교육 프로그램 운영이 어려운 실정임\n\n• 지역 문화시설(박물관, 미술관, 도서관)과의 연계가 미흡하여 다양한 문화 경험 기회가 제한됨", "improvement": "• 2025년부터 문화 담당자가 월 1회 이상 외부 문화체험 프로그램(연극 관람, 영화 감상, 미술 전시 관람 등)을 기획·운영하고, 연간 문화체험 횟수를 기존 4회에서 12회 이상으로 확대하여 아동들의 문화적 경험 기회를 대폭 늘림\n\n• 2025년 1분기까지 지역 예술 단체 및 음악·미술 전문 강사 2명과 외부 강사 협약을 체결하고, 주 2회 정기 예술 교육(음악 1회, 미술 1회)을 실시하여 전문적인 예술 교육 프로그램을 안정적으로 운영함\n\n• 2025년 상반기까지 인근 박물관 2개소, 미술관 1개소, 도서관 2개소와 문화 연계 MOU를 체결하고, 분기별 1회 이상 연계 프로그램(박물관 탐방, 도서관 독서교실 등)을 정례화하여 다양한 문화 경험 기회를 제공함"},
                {"area": "정서지원", "problem": "• 개별 상담 시간이 아동 1인당 월 1회 30분으로 부족하여 깊이 있는 정서 지원에 한계가 있음\n\n• 위기 아동(학대, 방임, 가정 문제) 조기 발견 및 개입 체계가 미흡하여 적시 지원이 이루어지지 못함\n\n• 또래 관계 갈등이나 학교폭력 피해 아동을 위한 전문 프로그램이 부재한 상황임", "improvement": "• 2025년 1분기까지 전문 상담사 1명을 추가 채용하여 상담 인력을 2명으로 확대하고, 개별 상담 시간을 기존 월 1회 30분에서 월 2회 50분으로 확대 운영하며, 상담 예약 시스템을 도입하여 상담 접근성을 대폭 강화함\n\n• 2025년 상반기까지 정서지원 담당자가 위기 아동 조기 발견을 위한 표준화된 스크리닝 도구를 도입하여 분기별 전수 평가를 실시하고, 아동보호전문기관과 긴급 핫라인을 구축하여 위기 상황 발생 시 24시간 이내 신속 개입이 가능한 체계를 마련함\n\n• 2025년 1분기부터 또래 관계 향상을 위한 사회성 증진 집단 프로그램(주 1회, 8주 과정)을 개발·운영하고, 학교폭력 피해 아동을 위한 전문 회복 프로그램(개별 상담 + 집단 치료)을 도입하여 정서적 안정을 지원함"},
                {"area": "지역사회연계", "problem": "• 자원봉사자 참여가 연간 50명 수준으로 저조하고 자원봉사 활동의 지속성이 낮아 안정적 운영이 어려움\n\n• 지역 내 기업, 단체와의 후원 및 협력 관계가 부족하여 프로그램 운영 재원 확보에 제한이 있음\n\n• 지역 유관기관(학교, 주민센터, 복지관)과의 정보 공유 및 협력 체계가 구축되지 않아 통합 지원에 한계가 있음", "improvement": "• 2025년 상반기까지 홍보 담당자가 자원봉사 모집 홍보를 대학교, 기업, 종교단체 등으로 확대하고, 월 1회 정기 봉사단(20명 규모)을 조직하여 운영하며, 봉사자 감사 행사를 연 2회 개최하여 연간 100명 이상의 봉사자가 지속적으로 참여하도록 유도함\n\n• 2025년 1분기까지 후원 담당자가 지역 내 주요 기업 10개소의 CSR 담당자 네트워크를 구축하고, 최소 5개 기업과 연간 후원 협약을 체결하여 프로그램 운영 재원 3,000만 원 이상을 안정적으로 확보함\n\n• 2025년 상반기까지 지역 유관기관(초등학교 3개교, 주민센터 2개소, 복지관 1개소) 실무자 협의체를 구성하고, 분기별 1회 사례회의를 정례화하며, 정보 공유 시스템을 구축하여 아동 통합 지원 체계를 마련함"}
            ],
            "total_review_table": [
                {"category": "운영평가", "content": "● **시설운영 현황**: 2024년 연간 사업 운영은 전반적으로 목표 달성률 85%를 기록하며 성공적으로 수행되었습니다. 시설 이용률은 평균 92%를 유지하였으며, 월별 운영 현황을 분석한 결과 상반기 대비 하반기 이용률이 8% 증가하는 긍정적인 추세를 보였습니다.\n\n● **인력관리 성과**: 돌봄 교사 4명, 학습지도사 2명, 상담사 1명 등 총 7명의 인력이 안정적으로 근무하였으며, 직원 이직률은 0%로 인력 안정성이 확보되었습니다.\n\n● **예산집행 현황**: 예산 집행률은 98.5%로 계획된 예산을 효율적으로 사용하였으며, 특히 프로그램 운영비와 인건비 집행이 계획대로 이루어졌습니다.\n\n● **안전관리**: 시설 안전관리 측면에서는 월 1회 정기 안전점검을 실시하여 안전사고 발생 건수 0건을 기록하였습니다.\n\n● **개선필요 사항**: 저녁 돌봄 시간 연장에 대한 요구가 높아 차년도 운영 시간 조정이 필요한 것으로 파악되었습니다. 행정 업무의 효율화를 위해 전자문서 시스템을 도입하여 업무 처리 시간이 30% 단축되는 성과를 거두었습니다."},
                {"category": "아동평가", "content": "● **이용현황**: 2024년 이용 아동 수는 총 58명으로 전년 대비 15% 증가하였으며, 정원 대비 이용률 96.7%를 기록하였습니다. 아동 출석률은 평균 94.2%로 매우 높은 수준을 유지하였습니다.\n\n● **발달성과**: 아동 발달 평가 결과, 사회성 발달 영역에서 전년 대비 12% 향상된 결과를 보였으며, 특히 또래 관계 형성 및 협동심 영역에서 두드러진 성장이 관찰되었습니다. 학습 능력 측면에서는 이용 아동의 학업성취도가 평균 8% 향상되었습니다.\n\n● **정서발달**: 정서 발달 측면에서는 분기별 정서 평가 결과 전체 아동의 87%가 안정적인 정서 상태를 유지하고 있는 것으로 나타났습니다. 위기 아동 8명에 대해서는 개별 사례관리를 통해 5명이 안정화되는 성과를 거두었습니다.\n\n● **만족도**: 아동 만족도 조사 결과 평균 4.3점(5점 만점)을 기록하였으며, 특히 급식과 간식에 대한 만족도가 4.6점으로 가장 높게 나타났습니다. 향후 개별 아동의 특성을 고려한 맞춤형 지원 강화가 필요합니다."},
                {"category": "프로그램평가", "content": "● **프로그램 운영실적 분석**: 2024년 총 15개 프로그램을 운영하였으며, 프로그램별 참여율은 평균 89%를 기록하였습니다. 교육 프로그램(학습지도, 독서지도, 디지털교육)은 총 48회 운영되었으며, 참여 아동의 학습 능력 향상에 크게 기여하였습니다. 특히 디지털 리터러시 교육은 신규 도입된 프로그램으로 아동과 보호자 모두에게 높은 호응을 얻었으며, 참여율 95%와 만족도 4.7점(5점 만점)을 달성하여 차년도 확대 운영을 결정하였습니다.\n\n● **문화프로그램 성과 및 과제**: 문화 프로그램(체험활동, 예술교육)은 연간 12회 운영되었으나, 다양성 측면에서 개선이 필요한 것으로 평가되었습니다. 박물관 탐방, 미술관 견학 등 외부 체험활동에 대한 아동들의 만족도가 4.2점으로 높게 나타났으나, 체험 횟수가 연간 4회에 불과하여 차년도에는 월 1회 이상으로 확대 운영할 계획입니다.\n\n● **정서지원 프로그램 효과성**: 정서지원 프로그램(개별상담, 집단상담, 또래활동)은 총 156회 운영되었으며, 참여 아동의 정서 안정에 긍정적인 효과를 보였습니다. 특히 사회성 향상 집단 프로그램 참여 아동 20명 중 18명(90%)이 또래관계 개선을 보고하였으며, 정서 불안 지수가 평균 25% 감소하는 성과를 거두었습니다."},
                {"category": "후원활동측면", "content": "● **후원현황**: 2024년 후원금 모금액은 총 2,500만 원으로 전년 대비 20% 증가하였습니다. 정기후원자 수는 45명으로 10명이 신규 등록되었으며, 후원자 유지율은 92%를 기록하였습니다.\n\n● **기업연계**: 기업 후원은 지역 내 5개 기업과 협약을 체결하여 연간 1,000만 원 상당의 현금 및 현물 후원을 확보하였습니다. 후원금은 프로그램 운영비(40%), 아동 지원비(35%), 시설 운영비(25%)로 투명하게 집행되었습니다.\n\n● **자원봉사**: 자원봉사자 참여는 연간 62명으로 목표(50명) 대비 124% 달성하였습니다. 봉사 분야는 학습지도(40%), 급식보조(30%), 프로그램 보조(30%)로 다양하게 참여하였습니다.\n\n● **지역네트워크**: 지역사회 연계 활동으로는 인근 초등학교 3개교와 협력하여 방과후 연계 프로그램을 운영하였으며, 주민센터, 도서관 등 5개 유관기관과 네트워크를 구축하였습니다."},
                {"category": "환류방안", "content": "● **운영시간 개선 계획**: 운영 시간 연장 요구에 대응하여 저녁 돌봄 시간을 18시에서 20시로 연장하고, 이를 위한 야간 돌봄 인력 1명을 추가 배치합니다. 2024년 학부모 설문조사 결과 전체 응답자의 67%가 운영 시간 연장을 희망하였으며, 특히 맞벌이 가정의 경우 87%가 20시 이후까지 돌봄이 필요하다고 응답하였습니다. 이에 2025년 3월부터 야간 돌봄 시간을 단계적으로 연장하여 운영할 예정입니다.\n\n● **문화프로그램 확대 방안**: 문화체험 프로그램 다양성 부족 문제 해결을 위해 월 1회 이상 외부 문화체험 활동을 실시하고, 지역 문화시설과 연계 협약을 체결합니다. 2025년 1분기까지 지역 내 박물관 2개소, 미술관 1개소, 도서관 3개소와 MOU를 체결하고, 분기별 특별 문화체험 프로그램을 기획하여 아동들의 다양한 문화 경험 기회를 확대할 계획입니다.\n\n● **맞춤형 학습지원 강화 방안**: 개별 아동 맞춤형 지원 강화를 위해 학습 진단 도구를 도입하고 개별 학습 계획(ILP)을 수립하여 체계적인 학습 지원을 제공합니다. 2025년 상반기까지 표준화된 학습 진단 도구를 도입하고, 아동별 학습 수준과 특성에 맞는 개별화 교육 계획을 수립하여 학업성취도 향상을 지원합니다.\n\n● **위기아동 지원체계 구축 방안**: 위기 아동 조기 발견 체계 구축을 위해 분기별 스크리닝을 실시하고 아동보호전문기관과 협력 체계를 강화합니다. 2025년부터 전문 상담사가 분기별 정서 스크리닝을 실시하고, 위기 징후가 발견된 아동에 대해서는 48시간 이내 개입하는 신속대응 체계를 마련합니다."}
            ],
            "satisfaction_survey": {
                "total_respondents": 30,
                "survey_data": [
                    {"문항": "1. 급식 및 간식의 질과 위생 상태", "5점": 15, "4점": 10, "3점": 3, "2점": 1, "1점": 1},
                    {"문항": "2. 프로그램 내용과 운영", "5점": 14, "4점": 11, "3점": 3, "2점": 1, "1점": 1},
                    {"문항": "3. 시설의 안전 및 위생 관리", "5점": 16, "4점": 9, "3점": 3, "2점": 1, "1점": 1},
                    {"문항": "4. 담당 선생님의 아동 지도 능력", "5점": 15, "4점": 10, "3점": 3, "2점": 1, "1점": 1},
                    {"문항": "5. 학습지도 프로그램의 효과", "5점": 12, "4점": 12, "3점": 4, "2점": 1, "1점": 1},
                    {"문항": "6. 아동의 정서적 지원", "5점": 13, "4점": 11, "3점": 4, "2점": 1, "1점": 1},
                    {"문항": "7. 부모 상담 및 소통", "5점": 11, "4점": 12, "3점": 5, "2점": 1, "1점": 1},
                    {"문항": "8. 문화체험 활동", "5점": 10, "4점": 11, "3점": 6, "2점": 2, "1점": 1},
                    {"문항": "9. 시설 운영 시간", "5점": 11, "4점": 11, "3점": 5, "2점": 2, "1점": 1},
                    {"문항": "10. 전반적인 시설 운영", "5점": 14, "4점": 11, "3점": 3, "2점": 1, "1점": 1}
                ],
                "subjective_question": "기타 건의사항 및 개선 의견",
                "subjective_analysis": "● **긍정적 의견**: 주관식 응답 분석 결과, 학부모들은 전반적으로 시설 운영에 높은 만족감을 표현하였습니다. 가장 많이 언급된 긍정적 의견으로는 담당 선생님들의 헌신적인 아동 지도 태도, 안전하고 청결한 시설 환경, 영양가 있는 급식 제공 등이 있었습니다.\n\n● **방역신뢰**: 특히 코로나19 이후 방역 관리에 대한 신뢰가 높게 나타났으며, 아동들이 즐겁게 센터에 다니고 있다는 피드백이 다수였습니다.\n\n● **개선요청**: 개선이 필요한 부분으로는 저녁 돌봄 시간 연장 요청이 가장 많았으며(전체 응답의 35%), 다양한 문화체험 프로그램 확대(28%), 학습지도 프로그램의 다양화(20%) 순으로 나타났습니다. 일부 학부모는 주말 프로그램 신설과 방학 기간 특별 프로그램을 희망하였습니다.",
                "overall_suggestion": "● **조사결과**: 2024년 만족도 조사 결과, 전체 평균 만족도 점수는 5점 만점에 4.29점으로 전년 대비 0.15점 상승하였습니다. 가장 높은 만족도를 보인 항목은 '급식 및 간식'(4.6점)과 '시설 안전 관리'(4.5점)입니다.\n\n● **개선필요**: 반면 '문화체험 활동'(3.9점)과 '운영 시간'(4.0점)은 상대적으로 낮은 점수를 기록하여 개선이 필요한 것으로 나타났습니다.\n\n● **제언사항**: 첫째 문화체험 프로그램을 월 1회 이상으로 확대하고 다양한 장르를 포함할 것을 권고합니다. 둘째, 맞벌이 가정의 요구를 반영하여 저녁 돌봄 시간을 20시까지 연장 운영하는 방안을 검토해야 합니다. 셋째, 학습지도의 효과성 향상을 위해 개별 맞춤형 학습 계획 수립과 정기적인 학습 성과 공유 시스템을 도입할 필요가 있습니다."
            },
            "purpose_text": "본 사업은 지역사회 내 돌봄이 필요한 만 18세 미만 아동, 특히 저소득층, 한부모가정, 다문화가정 등 취약계층 가정의 아동을 대상으로 가정의 돌봄 기능을 보완하고 사회적 안전망을 강화하여 모든 아동이 차별 없이 건강하게 성장할 수 있는 환경을 조성함을 목적으로 한다. 이를 위해 안전하고 위생적인 돌봄 환경에서 균형 잡힌 급식과 간식을 제공하고, 기초학습 지도 및 정서적 지원 프로그램을 통해 아동의 전인적 발달을 도모한다. 또한 가정-학교-지역사회 간 유기적 연계 체계를 구축하여 아동과 가정이 필요로 하는 통합적 지원 서비스를 제공하고, 아동의 권리를 존중하며 참여권을 보장하는 민주적 운영 원칙을 실현함으로써 지역사회 아동복지 향상에 기여하고자 한다.",
            "goals_text": "● **건강하고 안전한 권리 기반 돌봄 환경 조성**: 아동의 기본적인 생존권과 보호권을 보장하기 위해 일상생활 및 위생 관리 교육을 정례화하여 아동 스스로 건강한 생활습관을 형성하도록 지원한다. 균형 잡힌 영양 급식을 통해 성장기 아동의 신체 발달을 촉진하고, 안전하고 청결한 시설 환경을 유지하여 쾌적한 돌봄 공간을 제공한다. 아동 권리 헌장에 기반한 인권 교육을 분기별로 실시하고, 아동 자치회 운영을 통해 아동의 참여권과 의사표현권을 적극 보장한다.\n\n● **전인적 성장을 위한 맞춤형 교육 지원**: 기초 학습 지도를 넘어 문해력 향상을 위한 체계적인 독서 지도 프로그램을 운영하여 아동의 학습 능력 기반을 강화한다. 학습 진단 도구를 활용하여 개별 아동의 수준과 특성에 맞는 맞춤형 학습 계획을 수립하고, 정기적인 학습 성과 모니터링을 통해 교육 격차 해소에 기여한다. 예술(미술, 음악) 및 체육 활동을 통해 아동의 창의성과 신체 발달을 균형 있게 지원하며, 긍정적 강화 기법을 활용한 학습 동기 부여로 자기주도적 학습 태도를 함양한다.\n\n● **다양한 문화체험을 통한 사회성 함양**: 연간 계획에 따른 체계적인 문화체험 활동(박물관, 미술관, 공연 관람 등)을 통해 아동의 문화적 감수성과 심미적 안목을 키운다. 계절별 캠프 및 야외 체험활동을 실시하여 또래 간 협동심과 사회성을 발달시키고, 일상에서 벗어난 다양한 경험을 통해 정서적 해소와 심리적 안정감을 제공한다. 지역사회 문화시설과의 연계를 강화하여 아동들이 지역 문화자원을 적극 활용할 수 있도록 지원한다.\n\n● **정서적 안정과 가정기능 회복 지원**: 전문 상담사에 의한 개별 심리상담 및 집단 상담 프로그램을 운영하여 아동의 정서적 안정과 심리적 건강을 지원한다. 체계적인 사례관리를 통해 위기 아동을 조기에 발견하고 신속하게 개입하며, 아동보호전문기관 등 유관기관과의 협력 체계를 강화한다. 부모 상담 및 가족 지원 프로그램을 통해 가정의 양육 기능을 보완하고, 가족 관계 회복을 위한 통합적 지원 서비스를 제공한다.\n\n● **지속 가능한 운영과 지역사회 연계 강화**: 자원봉사자 및 후원자 관리 체계를 고도화하여 안정적인 인적·물적 자원을 확보하고, 후원금 사용의 투명성을 확보하기 위해 분기별 사용 보고서를 발행한다. 지역 내 학교, 주민센터, 복지관 등 유관기관과의 네트워크를 구축하여 아동에 대한 통합적 지원 체계를 마련하고, 정기적인 실무자 협의체 운영을 통해 정보 공유와 협력을 강화한다. 사업 운영의 지속가능성을 위해 재정 건전성을 유지하고 지역사회 참여를 확대한다."
        },
        "part2_programs": {
            "보호": {
                "subcategories": ["생활", "안전", "가족기능강화"],
                "detail_table": [
                    {"sub_area": "생활", "program_name": "위생관리 지도", "expected_effect": "아동들이 올바른 위생 습관을 형성하고 청결한 생활 환경을 유지하여 건강한 일상생활을 영위할 수 있습니다.", "target": "전체 아동", "count": "50명", "cycle": "매일", "content": "● **개인위생 점검**: 손씻기, 양치질 등 개인위생 습관 지도 및 점검표 관리\n● **환경위생 관리**: 시설 내 청소 및 정리정돈 습관 형성 지도"},
                    {"sub_area": "생활", "program_name": "영양 급식", "expected_effect": "균형 잡힌 영양 섭취를 통해 성장기 아동의 신체 발달을 촉진하고 건강한 식습관을 형성합니다.", "target": "전체 아동", "count": "50명", "cycle": "주 5회", "content": "● **급식 운영**: 영양사 검수 식단에 따른 급식 및 간식 제공\n● **식습관 교육**: 올바른 식사 예절 및 편식 교정 지도"},
                    {"sub_area": "안전", "program_name": "야간 보호", "expected_effect": "맞벌이 가정 아동의 돌봄 공백을 해소하고 안전한 귀가를 보장하여 부모의 양육 부담을 경감합니다.", "target": "야간돌봄 아동", "count": "20명", "cycle": "주 5회", "content": "● **석식 제공**: 야간 돌봄 아동 대상 영양 석식 제공\n● **안전 귀가 지원**: 차량 운행 및 보호자 인수인계 확인"},
                    {"sub_area": "안전", "program_name": "안전 교육", "expected_effect": "다양한 위기 상황에 대한 대처 능력을 함양하여 아동의 안전 의식을 고취하고 사고를 예방합니다.", "target": "전체 아동", "count": "50명", "cycle": "월 1회", "content": "● **소방 대피 훈련**: 월 1회 관할 소방서 연계 화재 대피 훈련 실시\n● **교통 안전 교육**: 등하교 시 교통 안전 수칙 교육"},
                    {"sub_area": "가족기능강화", "program_name": "부모 교육", "expected_effect": "보호자의 양육 역량을 강화하여 가정 내 건강한 양육 환경을 조성하고 부모-자녀 관계를 개선합니다.", "target": "보호자", "count": "40명", "cycle": "분기 1회", "content": "● **전문가 초빙 강의**: 아동 발달 단계에 따른 양육 방법 교육\n● **양육 스트레스 관리**: 부모 자조 모임 및 상담 연계"}
                ],
                "eval_table": [
                    {"sub_area": "생활", "program_name": "위생관리 지도", "expected_effect": "● **위생습관 형성**: 참여 아동의 95% 이상이 올바른 손씻기 방법을 습득하여 일상에서 자발적으로 실천함\n● **질병예방 효과**: 청결 습관 정착으로 계절성 감염 질환 발생률이 전년 대비 30% 감소함", "main_plan": "매일 등원 시 개인위생 점검 및 월 1회 위생 교육 실시", "eval_method": "프로그램 일지, 관찰일지"},
                    {"sub_area": "생활", "program_name": "영양 급식", "expected_effect": "● **영양상태 개선**: 균형 잡힌 식단 제공으로 성장기 아동의 신체 발달 지표가 평균 10% 향상됨\n● **식습관 교정**: 편식 아동 비율이 40%에서 15%로 감소하여 건강한 식습관이 형성됨", "main_plan": "영양사 검수 식단에 따라 주 5회 급식 제공 및 식사 예절 지도", "eval_method": "프로그램 일지, 만족도 조사"},
                    {"sub_area": "안전", "program_name": "야간 보호", "expected_effect": "● **돌봄공백 해소**: 맞벌이 가정 아동 100%가 안전한 환경에서 저녁 시간을 보내며 돌봄 공백 해소\n● **안전귀가 달성**: 보호자 인수인계 완료율 100% 달성으로 안전 귀가 보장", "main_plan": "17시~20시 야간 돌봄 운영 및 차량 귀가 지원", "eval_method": "출석부, 관찰일지"},
                    {"sub_area": "안전", "program_name": "안전 교육", "expected_effect": "● **위기대응 역량 강화**: 참여 아동 100%가 화재 대피 훈련에서 3분 이내 대피를 완료함\n● **안전의식 향상**: 일상생활에서 안전 수칙 준수율이 20% 향상되어 사고 예방에 기여함", "main_plan": "월 1회 소방서 연계 화재대피 훈련 및 교통안전 교육 실시", "eval_method": "프로그램 일지, 관찰일지"},
                    {"sub_area": "가족기능강화", "program_name": "부모 교육", "expected_effect": "● **양육역량 강화**: 참여 보호자의 90%가 양육 스트레스 감소 및 자녀 이해도 향상을 보고함\n● **가정환경 개선**: 교육 참여 가정의 부모-자녀 관계 만족도가 평균 25% 향상됨", "main_plan": "분기별 전문가 초빙 양육 교육 및 부모 자조 모임 운영", "eval_method": "만족도 조사, 사전사후 검사"}
                ]
            },
            "교육": {
                "subcategories": ["성장과권리", "학습", "특기적성"],
                "detail_table": [
                    {"sub_area": "성장과권리", "program_name": "아동권리 교육", "expected_effect": "아동이 자신의 권리를 이해하고 존중받는 환경에서 성장하여 자존감과 시민 의식을 함양합니다.", "target": "전체 아동", "count": "50명", "cycle": "분기 1회", "content": "● **권리 교육**: UN 아동권리협약 기반 권리 교육 실시\n● **권리 캠페인**: 아동 권리 포스터 제작 및 전시"},
                    {"sub_area": "성장과권리", "program_name": "아동자치회의", "expected_effect": "민주적 의사결정 과정을 경험하고 자기 표현 능력을 향상시켜 참여권을 보장합니다.", "target": "대표 아동", "count": "10명", "cycle": "월 1회", "content": "● **자치회의 운영**: 월 1회 아동 대표 회의 진행\n● **의견 반영**: 아동 의견 시설 운영에 적극 반영"},
                    {"sub_area": "학습", "program_name": "기초학력 향상", "expected_effect": "기초 학력 보충을 통해 학습 격차를 해소하고 학업 성취도를 향상시킵니다.", "target": "초등학생", "count": "30명", "cycle": "주 3회", "content": "● **교과 학습 지도**: 국어, 수학 기초학력 보충 지도\n● **개별 맞춤 지도**: 학습 수준별 소그룹 지도"},
                    {"sub_area": "학습", "program_name": "독서 지도", "expected_effect": "독서 습관 형성을 통해 문해력과 사고력을 향상시키고 평생 학습의 기초를 마련합니다.", "target": "전체 아동", "count": "40명", "cycle": "주 1회", "content": "● **함께 읽기**: 연령별 도서 선정 및 독서 활동\n● **독후 활동**: 독서토론, 독후감 작성, 독서퀴즈"},
                    {"sub_area": "특기적성", "program_name": "피아노 교실", "expected_effect": "음악적 감수성과 표현력을 키우고 집중력 향상을 통해 정서적 안정에 기여합니다.", "target": "희망 아동", "count": "15명", "cycle": "주 2회", "content": "● **기초 이론 교육**: 악보 읽기 및 기본 이론 학습\n● **실기 지도**: 개인별 수준에 맞는 연주 지도"},
                    {"sub_area": "특기적성", "program_name": "축구 교실", "expected_effect": "신체 활동을 통해 건강한 체력을 기르고 팀 활동을 통해 협동심과 사회성을 함양합니다.", "target": "희망 아동", "count": "20명", "cycle": "주 1회", "content": "● **기초 체력 훈련**: 준비운동 및 기초 체력 강화\n● **축구 기술 지도**: 패스, 슛, 드리블 등 기술 훈련"}
                ],
                "eval_table": [
                    {"sub_area": "성장과권리", "program_name": "아동권리 교육", "expected_effect": "● **권리인식 향상**: 참여 아동의 85% 이상이 자신의 권리 4가지 이상을 정확하게 인지하고 설명할 수 있음\n● **자존감 증진**: 권리 교육 후 자아존중감 척도 점수가 평균 15% 향상됨", "main_plan": "분기별 UN 아동권리협약 기반 권리 교육 및 캠페인 활동 전개", "eval_method": "사전사후 검사, 관찰일지"},
                    {"sub_area": "성장과권리", "program_name": "아동자치회의", "expected_effect": "● **참여권 실현**: 아동 제안 사항의 70% 이상이 시설 운영에 실제 반영되어 민주적 참여 경험 제공\n● **리더십 함양**: 자치회 참여 아동의 자기표현 능력 및 의사결정 능력이 30% 향상됨", "main_plan": "월 1회 정기 자치회의 개최 및 의견 수렴 체계 운영", "eval_method": "프로그램 일지, 관찰일지"},
                    {"sub_area": "학습", "program_name": "기초학력 향상", "expected_effect": "● **학력향상 달성**: 참여 아동의 평균 학업성취도가 학기 초 대비 15% 향상됨\n● **학습격차 해소**: 기초학력 미달 아동 비율이 25%에서 10%로 감소함", "main_plan": "주 3회 국어, 수학 교과 보충 지도 및 수준별 소그룹 학습 운영", "eval_method": "사전사후 검사, 프로그램 일지"},
                    {"sub_area": "학습", "program_name": "독서 지도", "expected_effect": "● **독서습관 형성**: 참여 아동의 월평균 독서량이 2권에서 5권으로 150% 증가함\n● **문해력 향상**: 독해력 평가 점수가 평균 20% 향상되어 학습 기초 역량 강화", "main_plan": "주 1회 연령별 독서 활동 및 독서토론, 독후감 작성 지도", "eval_method": "프로그램 일지, 관찰일지"},
                    {"sub_area": "특기적성", "program_name": "피아노 교실", "expected_effect": "● **음악역량 발달**: 참여 아동 80% 이상이 기초 악보를 읽고 2곡 이상 연주할 수 있게 됨\n● **정서안정 효과**: 음악 활동을 통해 정서 안정 및 집중력이 25% 향상됨", "main_plan": "주 2회 개인별 수준 맞춤 피아노 실기 지도 및 연 2회 발표회 개최", "eval_method": "프로그램 일지, 만족도 조사"},
                    {"sub_area": "특기적성", "program_name": "축구 교실", "expected_effect": "● **체력증진 달성**: 참여 아동의 기초체력(지구력, 민첩성) 측정 결과가 평균 20% 향상됨\n● **사회성 발달**: 팀 활동을 통해 협동심 및 또래관계 만족도가 30% 향상됨", "main_plan": "주 1회 전문 코치 지도 하에 축구 기술 훈련 및 친선 경기 진행", "eval_method": "프로그램 일지, 관찰일지"}
                ]
            },
            "문화": {
                "subcategories": ["체험활동"],
                "detail_table": [
                    {"sub_area": "체험활동", "program_name": "박물관 탐방", "expected_effect": "역사와 문화에 대한 이해를 넓히고 다양한 체험을 통해 학습 동기와 호기심을 자극합니다.", "target": "전체 아동", "count": "50명", "cycle": "분기 1회", "content": "● **현장 학습**: 지역 박물관 및 역사 유적지 견학\n● **체험 활동**: 전통 문화 체험 및 만들기 활동"},
                    {"sub_area": "체험활동", "program_name": "여름/겨울 캠프", "expected_effect": "일상을 벗어난 단체 생활 경험을 통해 자립심과 사회성을 키우고 추억을 형성합니다.", "target": "전체 아동", "count": "40명", "cycle": "연 2회", "content": "● **캠프 활동**: 1박 2일 캠프 운영(자연체험, 레크리에이션)\n● **공동 생활**: 협동 미션 및 단체 활동"},
                    {"sub_area": "체험활동", "program_name": "공연 관람", "expected_effect": "문화 예술을 직접 체험하여 감수성과 심미적 안목을 키우고 정서적 풍요를 경험합니다.", "target": "전체 아동", "count": "50명", "cycle": "분기 1회", "content": "● **공연 관람**: 연극, 뮤지컬, 영화 등 문화 공연 관람\n● **사후 활동**: 감상문 작성 및 토론"}
                ],
                "eval_table": [
                    {"sub_area": "체험활동", "program_name": "박물관 탐방", "expected_effect": "● **문화이해 확장**: 참여 아동의 90% 이상이 체험한 역사/문화 내용을 정확히 설명할 수 있음\n● **학습동기 향상**: 현장 학습 후 관련 교과 흥미도가 25% 향상되어 능동적 학습 태도 형성", "main_plan": "분기별 지역 박물관 및 역사 유적지 견학 프로그램 운영", "eval_method": "프로그램 일지, 만족도 조사"},
                    {"sub_area": "체험활동", "program_name": "여름/겨울 캠프", "expected_effect": "● **자립심 강화**: 단체 생활 경험을 통해 참여 아동의 85%가 자기관리 능력 향상을 보임\n● **사회성 증진**: 협동 활동을 통해 또래관계 만족도가 35% 향상되고 추억 형성", "main_plan": "연 2회(하계/동계) 1박 2일 캠프 운영 및 다양한 체험 활동 제공", "eval_method": "만족도 조사, 관찰일지"},
                    {"sub_area": "체험활동", "program_name": "공연 관람", "expected_effect": "● **문화감수성 발달**: 참여 아동의 문화예술 관심도가 40% 향상되어 심미적 안목 형성\n● **정서적 풍요**: 공연 관람 후 정서 안정 및 스트레스 해소 효과로 긍정정서 20% 증가", "main_plan": "분기별 연극, 뮤지컬, 영화 등 문화 공연 관람 기회 제공", "eval_method": "프로그램 일지, 만족도 조사"}
                ]
            },
            "정서지원": {
                "subcategories": ["상담"],
                "detail_table": [
                    {"sub_area": "상담", "program_name": "개별 심리상담", "expected_effect": "아동 개인의 심리적 어려움을 파악하고 전문적인 상담을 통해 정서적 안정을 도모합니다.", "target": "전체 아동", "count": "50명", "cycle": "수시", "content": "● **정기 상담**: 전문 상담사와 월 1회 이상 개별 상담\n● **위기 개입**: 위기 상황 발생 시 즉시 상담 지원"},
                    {"sub_area": "상담", "program_name": "집단 상담", "expected_effect": "또래 관계 속에서 사회성을 향상시키고 공감 능력과 의사소통 기술을 발달시킵니다.", "target": "희망 아동", "count": "20명", "cycle": "주 1회", "content": "● **사회성 프로그램**: 8주 과정 사회성 향상 집단 프로그램\n● **또래 활동**: 협동 게임 및 역할극 활동"},
                    {"sub_area": "상담", "program_name": "보호자 상담", "expected_effect": "보호자와의 소통을 강화하여 가정-시설 간 연계를 촉진하고 양육 환경을 개선합니다.", "target": "보호자", "count": "40명", "cycle": "분기 1회", "content": "● **정기 상담**: 분기별 보호자 상담 실시\n● **가정 연계**: 가정 내 아동 지원 방안 협의"}
                ],
                "eval_table": [
                    {"sub_area": "상담", "program_name": "개별 심리상담", "expected_effect": "● **정서안정 달성**: 상담 참여 아동의 정서불안 지수가 평균 30% 감소함\n● **문제해결 향상**: 위기 아동 대상 개입 후 80%가 안정화되어 적응력 향상", "main_plan": "전문 상담사의 월 2회 개별 상담 및 위기 상황 시 즉각 개입 체계 운영", "eval_method": "프로그램 일지, 관찰일지"},
                    {"sub_area": "상담", "program_name": "집단 상담", "expected_effect": "● **사회성 증진**: 프로그램 참여 아동의 90%가 또래관계 개선을 경험함\n● **공감능력 향상**: 사회성 척도 검사 결과 공감 능력 점수가 평균 25% 향상됨", "main_plan": "8주 과정 사회성 향상 집단 프로그램 운영 및 또래 활동 지원", "eval_method": "사전사후 검사, 관찰일지"},
                    {"sub_area": "상담", "program_name": "보호자 상담", "expected_effect": "● **양육소통 강화**: 가정-시설 간 정보 공유로 아동 지원 일관성 90% 달성\n● **가정환경 개선**: 상담 참여 가정의 양육 환경 만족도가 20% 향상됨", "main_plan": "분기별 정기 보호자 상담 실시 및 가정 연계 지원 방안 협의", "eval_method": "프로그램 일지, 만족도 조사"}
                ]
            },
            "지역사회연계": {
                "subcategories": ["연계"],
                "detail_table": [
                    {"sub_area": "연계", "program_name": "자원봉사 관리", "expected_effect": "지역사회 자원봉사 인력을 활용하여 다양한 프로그램을 운영하고 아동에게 긍정적 역할 모델을 제공합니다.", "target": "전체 아동", "count": "50명", "cycle": "상시", "content": "● **봉사자 모집**: 대학생, 지역 주민 자원봉사자 모집 관리\n● **활동 배치**: 학습지도, 급식보조, 프로그램 보조 배치"},
                    {"sub_area": "연계", "program_name": "후원 연계", "expected_effect": "안정적인 후원 자원을 확보하여 시설 운영 및 프로그램 질 향상에 기여합니다.", "target": "시설", "count": "-", "cycle": "상시", "content": "● **후원자 관리**: 정기 후원자 감사 행사 및 소식지 발송\n● **기업 연계**: 지역 기업 CSR 연계 후원 확보"},
                    {"sub_area": "연계", "program_name": "기관 협력", "expected_effect": "유관 기관과의 네트워크를 통해 아동에게 통합적 지원 서비스를 제공합니다.", "target": "시설", "count": "-", "cycle": "분기 1회", "content": "● **실무자 협의체**: 학교, 주민센터, 복지관 등 분기별 협의\n● **사례 공유**: 아동 지원 사례 공유 및 연계"}
                ],
                "eval_table": [
                    {"sub_area": "연계", "program_name": "자원봉사 관리", "expected_effect": "● **인적자원 확보**: 연간 봉사자 100명 이상 확보로 안정적 프로그램 운영 지원\n● **역할모델 제공**: 대학생 봉사자와의 교류를 통해 아동의 진로 인식 및 학습 동기 20% 향상", "main_plan": "연중 자원봉사자 모집 및 월 1회 정기 봉사단 운영, 감사 행사 개최", "eval_method": "프로그램 일지, 만족도 조사"},
                    {"sub_area": "연계", "program_name": "후원 연계", "expected_effect": "● **재정안정성 확보**: 연간 후원금 3,000만 원 이상 확보로 프로그램 운영 재원 안정화\n● **후원자 유지**: 정기 후원자 유지율 90% 이상 달성으로 지속가능한 운영 기반 구축", "main_plan": "정기 후원자 감사 행사 및 소식지 발송, 기업 CSR 연계 활동 추진", "eval_method": "프로그램 일지, 관찰일지"},
                    {"sub_area": "연계", "program_name": "기관 협력", "expected_effect": "● **통합지원 실현**: 유관기관 연계를 통해 아동 통합 사례관리 대상의 85%가 개선 성과 달성\n● **네트워크 강화**: 분기별 협의체 운영으로 정보 공유 및 협력 사업 5건 이상 추진", "main_plan": "학교, 주민센터, 복지관 등 유관기관 협의체 분기별 운영 및 사례 공유", "eval_method": "프로그램 일지, 관찰일지"}
                ]
            }
        },
        "part3_monthly_plan": {
            "1월": [
                {"big_category": "보호", "mid_category": "생활", "program_name": "위생관리 지도", "target": "전체 아동", "staff": "돌봄교사", "content": "● **개인위생 점검**: 손씻기, 양치질 등 개인위생 습관 지도 및 점검표 관리\n● **환경위생 관리**: 시설 내 청소 및 정리정돈 습관 형성 지도"},
                {"big_category": "보호", "mid_category": "생활", "program_name": "영양 급식", "target": "전체 아동", "staff": "돌봄교사", "content": "● **급식 운영**: 영양사 검수 식단에 따른 급식 및 간식 제공\n● **식습관 교육**: 올바른 식사 예절 및 편식 교정 지도"},
                {"big_category": "보호", "mid_category": "안전", "program_name": "안전 교육", "target": "전체 아동", "staff": "사회복지사", "content": "● **동계 안전교육**: 빙판길 보행, 난방기구 사용 안전 교육\n● **대피 훈련**: 화재대피 훈련 실시"},
                {"big_category": "교육", "mid_category": "학습", "program_name": "기초학력 향상", "target": "초등학생", "staff": "학습지도사", "content": "● **국어, 수학 지도**: 기초학력 보충 지도\n● **개별 맞춤 지도**: 학습 수준별 소그룹 지도"},
                {"big_category": "문화", "mid_category": "체험활동", "program_name": "신년맞이 행사", "target": "전체 아동", "staff": "사회복지사", "content": "● **새해 소망 나누기**: 새해 목표 설정 및 다짐\n● **겨울방학 특별 프로그램**: 특기적성 집중 프로그램 운영"}
            ],
            "2월": [
                {"big_category": "보호", "mid_category": "생활", "program_name": "위생관리 지도", "target": "전체 아동", "staff": "돌봄교사", "content": "● **개인위생 점검**: 손씻기, 양치질 등 개인위생 습관 지도\n● **환경위생 관리**: 시설 내 청소 및 정리정돈"},
                {"big_category": "보호", "mid_category": "생활", "program_name": "영양 급식", "target": "전체 아동", "staff": "돌봄교사", "content": "● **급식 운영**: 균형 잡힌 영양 식단 제공\n● **식습관 교육**: 올바른 식사 예절 지도"},
                {"big_category": "교육", "mid_category": "학습", "program_name": "기초학력 향상", "target": "초등학생", "staff": "학습지도사", "content": "● **학년말 정리**: 학기 중 학습 내용 복습\n● **새학기 준비**: 신학년 준비 학습"},
                {"big_category": "문화", "mid_category": "체험활동", "program_name": "설날 전통문화 체험", "target": "전체 아동", "staff": "사회복지사", "content": "● **전통놀이**: 윷놀이, 제기차기 등 민속놀이 체험\n● **한복 체험**: 한복 입기 및 세배 예절 교육"}
            ],
            "3월": [
                {"big_category": "보호", "mid_category": "생활", "program_name": "위생관리 지도", "target": "전체 아동", "staff": "돌봄교사", "content": "● **개인위생 점검**: 손씻기, 양치질 등 개인위생 습관 지도\n● **환경위생 관리**: 시설 내 청소 및 정리정돈"},
                {"big_category": "보호", "mid_category": "생활", "program_name": "영양 급식", "target": "전체 아동", "staff": "돌봄교사", "content": "● **급식 운영**: 균형 잡힌 영양 식단 제공\n● **식습관 교육**: 올바른 식사 예절 지도"},
                {"big_category": "보호", "mid_category": "안전", "program_name": "안전 교육", "target": "전체 아동", "staff": "사회복지사", "content": "● **교통안전 교육**: 등하교 시 교통안전 수칙 교육\n● **대피 훈련**: 화재대피 훈련 실시"},
                {"big_category": "교육", "mid_category": "학습", "program_name": "기초학력 향상", "target": "초등학생", "staff": "학습지도사", "content": "● **새학기 적응 학습**: 신학년 교과 내용 예습\n● **학습 진단 평가**: 학기 초 학습 수준 진단"},
                {"big_category": "교육", "mid_category": "성장과권리", "program_name": "아동권리 교육", "target": "전체 아동", "staff": "사회복지사", "content": "● **권리 교육**: UN 아동권리협약 기반 교육\n● **자치회 구성**: 신학기 아동자치회 선거 진행"},
                {"big_category": "문화", "mid_category": "체험활동", "program_name": "박물관 탐방", "target": "전체 아동", "staff": "사회복지사", "content": "● **현장 학습**: 지역 역사박물관 견학\n● **체험 활동**: 전통 문화 체험 프로그램 참여"}
            ],
            "4월": [
                {"big_category": "보호", "mid_category": "생활", "program_name": "위생관리 지도", "target": "전체 아동", "staff": "돌봄교사", "content": "● **개인위생 점검**: 손씻기, 양치질 등 개인위생 습관 지도\n● **환경위생 관리**: 시설 내 청소 및 정리정돈"},
                {"big_category": "보호", "mid_category": "생활", "program_name": "영양 급식", "target": "전체 아동", "staff": "돌봄교사", "content": "● **급식 운영**: 균형 잡힌 영양 식단 제공\n● **식습관 교육**: 올바른 식사 예절 지도"},
                {"big_category": "교육", "mid_category": "학습", "program_name": "기초학력 향상", "target": "초등학생", "staff": "학습지도사", "content": "● **교과 학습 지도**: 국어, 수학 기초학력 보충\n● **개별 맞춤 지도**: 학습 수준별 소그룹 지도"},
                {"big_category": "교육", "mid_category": "특기적성", "program_name": "피아노 교실", "target": "희망 아동", "staff": "외부강사", "content": "● **기초 이론 교육**: 악보 읽기 및 기본 이론 학습\n● **실기 지도**: 개인별 수준에 맞는 연주 지도"},
                {"big_category": "문화", "mid_category": "체험활동", "program_name": "봄 환경정화 활동", "target": "전체 아동", "staff": "사회복지사", "content": "● **환경정화**: 지역사회 환경정화 활동 참여\n● **식목일 행사**: 나무 심기 체험 활동"}
            ],
            "5월": [
                {"big_category": "보호", "mid_category": "생활", "program_name": "위생관리 지도", "target": "전체 아동", "staff": "돌봄교사", "content": "● **개인위생 점검**: 손씻기, 양치질 등 개인위생 습관 지도\n● **환경위생 관리**: 시설 내 청소 및 정리정돈"},
                {"big_category": "보호", "mid_category": "생활", "program_name": "영양 급식", "target": "전체 아동", "staff": "돌봄교사", "content": "● **급식 운영**: 균형 잡힌 영양 식단 제공\n● **식습관 교육**: 올바른 식사 예절 지도"},
                {"big_category": "보호", "mid_category": "가족기능강화", "program_name": "부모 교육", "target": "보호자", "staff": "사회복지사", "content": "● **가정의 달 부모교육**: 양육 역량 강화 교육\n● **가족 참여 행사**: 가족과 함께하는 프로그램"},
                {"big_category": "교육", "mid_category": "학습", "program_name": "기초학력 향상", "target": "초등학생", "staff": "학습지도사", "content": "● **교과 학습 지도**: 국어, 수학 기초학력 보충\n● **개별 맞춤 지도**: 학습 수준별 소그룹 지도"},
                {"big_category": "문화", "mid_category": "체험활동", "program_name": "어린이날 행사", "target": "전체 아동", "staff": "사회복지사", "content": "● **어린이날 축하행사**: 레크리에이션 및 선물 증정\n● **놀이 안전 교육**: 안전한 놀이 방법 교육"}
            ],
            "6월": [
                {"big_category": "보호", "mid_category": "생활", "program_name": "위생관리 지도", "target": "전체 아동", "staff": "돌봄교사", "content": "● **개인위생 점검**: 손씻기, 양치질 등 개인위생 습관 지도\n● **환경위생 관리**: 시설 내 청소 및 정리정돈"},
                {"big_category": "보호", "mid_category": "생활", "program_name": "영양 급식", "target": "전체 아동", "staff": "돌봄교사", "content": "● **급식 운영**: 균형 잡힌 영양 식단 제공\n● **식습관 교육**: 올바른 식사 예절 지도"},
                {"big_category": "보호", "mid_category": "안전", "program_name": "안전 교육", "target": "전체 아동", "staff": "사회복지사", "content": "● **폭염 대비 교육**: 여름철 건강관리 및 열사병 예방\n● **수상 안전 교육**: 물놀이 안전 수칙 교육"},
                {"big_category": "교육", "mid_category": "학습", "program_name": "기초학력 향상", "target": "초등학생", "staff": "학습지도사", "content": "● **1학기 마무리 학습**: 학기 중 학습 내용 복습\n● **상반기 평가**: 학습 성과 점검 및 환류"},
                {"big_category": "문화", "mid_category": "체험활동", "program_name": "박물관 탐방", "target": "전체 아동", "staff": "사회복지사", "content": "● **현장 학습**: 과학관 또는 역사관 견학\n● **체험 활동**: 다양한 체험 프로그램 참여"}
            ],
            "7월": [
                {"big_category": "보호", "mid_category": "생활", "program_name": "위생관리 지도", "target": "전체 아동", "staff": "돌봄교사", "content": "● **개인위생 점검**: 손씻기, 양치질 등 개인위생 습관 지도\n● **환경위생 관리**: 여름철 위생 관리 강화"},
                {"big_category": "보호", "mid_category": "생활", "program_name": "영양 급식", "target": "전체 아동", "staff": "돌봄교사", "content": "● **급식 운영**: 균형 잡힌 영양 식단 제공\n● **식중독 예방**: 여름철 식품 위생 관리 강화"},
                {"big_category": "보호", "mid_category": "안전", "program_name": "안전 교육", "target": "전체 아동", "staff": "사회복지사", "content": "● **수상 안전 교육**: 물놀이 안전 수칙 교육\n● **폭염 대비 교육**: 열사병 예방 및 건강관리"},
                {"big_category": "교육", "mid_category": "학습", "program_name": "기초학력 향상", "target": "초등학생", "staff": "학습지도사", "content": "● **여름방학 학습**: 방학 중 자기주도 학습 지도\n● **독서 프로그램**: 여름 독서 캠프 운영"},
                {"big_category": "문화", "mid_category": "체험활동", "program_name": "여름 캠프", "target": "전체 아동", "staff": "사회복지사", "content": "● **1박 2일 캠프**: 자연체험 및 레크리에이션\n● **물놀이 활동**: 안전한 물놀이 프로그램 운영"}
            ],
            "8월": [
                {"big_category": "보호", "mid_category": "생활", "program_name": "위생관리 지도", "target": "전체 아동", "staff": "돌봄교사", "content": "● **개인위생 점검**: 손씻기, 양치질 등 개인위생 습관 지도\n● **환경위생 관리**: 여름철 위생 관리 강화"},
                {"big_category": "보호", "mid_category": "생활", "program_name": "영양 급식", "target": "전체 아동", "staff": "돌봄교사", "content": "● **급식 운영**: 균형 잡힌 영양 식단 제공\n● **식중독 예방**: 여름철 식품 위생 관리"},
                {"big_category": "교육", "mid_category": "학습", "program_name": "기초학력 향상", "target": "초등학생", "staff": "학습지도사", "content": "● **신학기 준비 학습**: 2학기 예습 및 복습\n● **독서 지도**: 여름 독서 마무리 및 독후 활동"},
                {"big_category": "문화", "mid_category": "체험활동", "program_name": "자연체험 활동", "target": "전체 아동", "staff": "사회복지사", "content": "● **생태 탐방**: 생태공원 또는 숲 체험 활동\n● **농촌 체험**: 농촌 일일 체험 프로그램"}
            ],
            "9월": [
                {"big_category": "보호", "mid_category": "생활", "program_name": "위생관리 지도", "target": "전체 아동", "staff": "돌봄교사", "content": "● **개인위생 점검**: 손씻기, 양치질 등 개인위생 습관 지도\n● **환경위생 관리**: 시설 내 청소 및 정리정돈"},
                {"big_category": "보호", "mid_category": "생활", "program_name": "영양 급식", "target": "전체 아동", "staff": "돌봄교사", "content": "● **급식 운영**: 균형 잡힌 영양 식단 제공\n● **식습관 교육**: 올바른 식사 예절 지도"},
                {"big_category": "보호", "mid_category": "안전", "program_name": "안전 교육", "target": "전체 아동", "staff": "사회복지사", "content": "● **교통안전 교육**: 등하교 시 교통안전 수칙 교육\n● **대피 훈련**: 화재대피 훈련 실시"},
                {"big_category": "교육", "mid_category": "학습", "program_name": "기초학력 향상", "target": "초등학생", "staff": "학습지도사", "content": "● **2학기 학습 지도**: 교과 내용 보충 학습\n● **개별 맞춤 지도**: 학습 수준별 소그룹 지도"},
                {"big_category": "문화", "mid_category": "체험활동", "program_name": "추석 전통문화 체험", "target": "전체 아동", "staff": "사회복지사", "content": "● **전통놀이**: 민속놀이 체험 활동\n● **송편 만들기**: 추석 전통 음식 만들기 체험"},
                {"big_category": "문화", "mid_category": "체험활동", "program_name": "박물관 탐방", "target": "전체 아동", "staff": "사회복지사", "content": "● **현장 학습**: 지역 박물관 견학\n● **체험 활동**: 역사 문화 체험 프로그램"}
            ],
            "10월": [
                {"big_category": "보호", "mid_category": "생활", "program_name": "위생관리 지도", "target": "전체 아동", "staff": "돌봄교사", "content": "● **개인위생 점검**: 손씻기, 양치질 등 개인위생 습관 지도\n● **환경위생 관리**: 시설 내 청소 및 정리정돈"},
                {"big_category": "보호", "mid_category": "생활", "program_name": "영양 급식", "target": "전체 아동", "staff": "돌봄교사", "content": "● **급식 운영**: 균형 잡힌 영양 식단 제공\n● **식습관 교육**: 올바른 식사 예절 지도"},
                {"big_category": "교육", "mid_category": "학습", "program_name": "기초학력 향상", "target": "초등학생", "staff": "학습지도사", "content": "● **교과 학습 지도**: 국어, 수학 기초학력 보충\n● **독서의 달 행사**: 독서 골든벨 및 독후감 대회"},
                {"big_category": "교육", "mid_category": "학습", "program_name": "독서 지도", "target": "전체 아동", "staff": "학습지도사", "content": "● **독서의 달 특별 프로그램**: 다독상, 독서왕 시상\n● **독서토론회**: 연령별 독서토론 활동"},
                {"big_category": "문화", "mid_category": "체험활동", "program_name": "가을 소풍", "target": "전체 아동", "staff": "사회복지사", "content": "● **야외 활동**: 자연 속 가을 소풍\n● **체험 활동**: 고구마 캐기 등 농촌 체험"}
            ],
            "11월": [
                {"big_category": "보호", "mid_category": "생활", "program_name": "위생관리 지도", "target": "전체 아동", "staff": "돌봄교사", "content": "● **개인위생 점검**: 손씻기, 양치질 등 개인위생 습관 지도\n● **환경위생 관리**: 시설 내 청소 및 정리정돈"},
                {"big_category": "보호", "mid_category": "생활", "program_name": "영양 급식", "target": "전체 아동", "staff": "돌봄교사", "content": "● **급식 운영**: 균형 잡힌 영양 식단 제공\n● **식습관 교육**: 올바른 식사 예절 지도"},
                {"big_category": "보호", "mid_category": "안전", "program_name": "안전 교육", "target": "전체 아동", "staff": "사회복지사", "content": "● **화재 예방 교육**: 겨울철 화재 예방 및 대피 훈련\n● **난방기구 안전 교육**: 안전한 난방기구 사용법"},
                {"big_category": "교육", "mid_category": "학습", "program_name": "기초학력 향상", "target": "초등학생", "staff": "학습지도사", "content": "● **학기말 정리 학습**: 2학기 학습 내용 복습\n● **개별 맞춤 지도**: 학습 부진 아동 특별 지도"},
                {"big_category": "지역사회연계", "mid_category": "연계", "program_name": "지역축제 참여", "target": "전체 아동", "staff": "사회복지사", "content": "● **지역 행사 참여**: 지역 문화축제 관람 및 참여\n● **봉사활동**: 지역사회 봉사활동 참여"}
            ],
            "12월": [
                {"big_category": "보호", "mid_category": "생활", "program_name": "위생관리 지도", "target": "전체 아동", "staff": "돌봄교사", "content": "● **개인위생 점검**: 손씻기, 양치질 등 개인위생 습관 지도\n● **환경위생 관리**: 시설 내 청소 및 정리정돈"},
                {"big_category": "보호", "mid_category": "생활", "program_name": "영양 급식", "target": "전체 아동", "staff": "돌봄교사", "content": "● **급식 운영**: 균형 잡힌 영양 식단 제공\n● **식습관 교육**: 올바른 식사 예절 지도"},
                {"big_category": "보호", "mid_category": "안전", "program_name": "안전 교육", "target": "전체 아동", "staff": "사회복지사", "content": "● **동계 안전교육**: 빙판길 보행, 난방기구 사용 안전\n● **화재 예방 교육**: 겨울철 화재 예방 교육"},
                {"big_category": "교육", "mid_category": "학습", "program_name": "기초학력 향상", "target": "초등학생", "staff": "학습지도사", "content": "● **연간 학습 정리**: 1년간 학습 내용 종합 복습\n● **겨울방학 학습 계획**: 방학 중 학습 계획 수립"},
                {"big_category": "문화", "mid_category": "체험활동", "program_name": "송년행사", "target": "전체 아동", "staff": "사회복지사", "content": "● **송년회**: 1년 감사 나눔 및 성과 발표회\n● **연말 시상식**: 우수 아동 및 봉사자 시상"},
                {"big_category": "문화", "mid_category": "체험활동", "program_name": "박물관 탐방", "target": "전체 아동", "staff": "사회복지사", "content": "● **현장 학습**: 연말 문화체험 활동\n● **공연 관람**: 크리스마스 공연 관람"}
            ]
        },
        "part4_budget_evaluation": {
            "budget_table": [
                {"category": "인건비", "amount": "120,000,000원", "details": "● **돌봄교사 4명**: 48,000,000원 (월 1,000,000원 × 4명 × 12개월)\n● **학습지도사 2명**: 36,000,000원 (월 1,500,000원 × 2명 × 12개월)\n● **상담사 1명**: 24,000,000원 (월 2,000,000원 × 1명 × 12개월)\n● **센터장 1명**: 12,000,000원 (월 1,000,000원 × 1명 × 12개월)"},
                {"category": "사업비", "amount": "35,000,000원", "details": "● **급식비**: 20,000,000원 (50명 × 2,000원 × 200일)\n● **프로그램 운영비**: 10,000,000원 (교육/문화 프로그램 운영)\n● **교재교구비**: 5,000,000원 (학습교재, 체육용품, 미술재료 등)"},
                {"category": "운영비", "amount": "15,000,000원", "details": "● **시설관리비**: 8,000,000원 (전기, 수도, 가스, 청소 등)\n● **통신비**: 2,000,000원 (인터넷, 전화 등)\n● **사무용품비**: 3,000,000원 (문구류, 소모품 등)\n● **보험료**: 2,000,000원 (아동 상해보험, 시설 화재보험)"},
                {"category": "체험활동비", "amount": "12,000,000원", "details": "● **현장학습비**: 6,000,000원 (분기별 박물관, 미술관 견학)\n● **캠프비**: 4,000,000원 (하계/동계 캠프 운영)\n● **공연관람비**: 2,000,000원 (분기별 문화공연 관람)"},
                {"category": "기타경비", "amount": "8,000,000원", "details": "● **차량유지비**: 4,000,000원 (셔틀버스 유류비, 정비비)\n● **행사비**: 2,000,000원 (어린이날, 송년행사 등)\n● **예비비**: 2,000,000원 (긴급 상황 대비)"},
                {"category": "총계", "amount": "190,000,000원", "details": "2025년 연간 예산 총액"}
            ],
            "feedback_summary": [
                {"area": "보호", "problem": "저녁 돌봄 시간 부족, 안전시설 노후화", "plan": "운영시간 20시까지 연장, CCTV 및 안전문 전면 교체"},
                {"area": "교육", "problem": "디지털 기기 부족, 교과 다양성 미흡", "plan": "태블릿 20대 구입, 영어/과학 전문강사 채용"},
                {"area": "문화", "problem": "문화체험 횟수 부족, 예술교육 미흡", "plan": "월 1회 이상 체험활동, 음악/미술 전문강사 확보"},
                {"area": "정서지원", "problem": "개별상담 시간 부족, 위기아동 개입 미흡", "plan": "상담사 1명 추가 채용, 위기아동 조기발견 체계 구축"},
                {"area": "지역사회연계", "problem": "자원봉사자 참여 저조, 후원 확보 미흡", "plan": "봉사자 모집 확대, 기업 CSR 연계 강화"}
            ]
        }
    }
