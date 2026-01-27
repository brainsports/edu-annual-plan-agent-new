import streamlit as st
import google.generativeai as genai
import os
from docx import Document
from io import BytesIO

# 1. 페이지 설정 및 제목 (선생님이 원하신 프로 버전 UI)
st.set_page_config(page_title="참참 3.1 프로 분석기", layout="wide")
st.title("🚀 참참 3.1 프로: 통합 편집 및 분석 시스템")
st.markdown("---")

# 2. API 설정 (현재 유료 키 반영)
API_KEY = os.environ.get("GEMINI_API_KEY")
genai.configure(api_key=API_KEY)
model = genai.GenerativeModel('gemini-1.5-flash')


# 3. 문서 저장 함수
def save_to_word(title, content):
    doc = Document()
    doc.add_heading(title, 0)
    doc.add_paragraph(content)
    bio = BytesIO()
    doc.save(bio)
    return bio.getvalue()


# 4. 분석 및 파트 분할 로직
def run_analysis(filename, text):
    prompt = f"""
    당신은 베테랑 사회복지사입니다. 다음 '{filename}'의 내용을 분석하여 
    반드시 [PART1], [PART2], [PART3], [PART4] 구분자를 넣어 작성하세요.
    [PART1] 성과 요약 (핵심 수치 중심)
    [PART2] 문제점 및 한계
    [PART3] 향후 개선 계획
    [PART4] 총평
    내용: {text[:8000]}
    """
    try:
        response = model.generate_content(prompt).text
        parts = {"P1": "", "P2": "", "P3": "", "P4": ""}
        if "[PART1]" in response:
            parts["P1"] = response.split("[PART1]")[1].split(
                "[PART2]")[0].strip()
        if "[PART2]" in response:
            parts["P2"] = response.split("[PART2]")[1].split(
                "[PART3]")[0].strip()
        if "[PART3]" in response:
            parts["P3"] = response.split("[PART3]")[1].split(
                "[PART4]")[0].strip()
        if "[PART4]" in response:
            parts["P4"] = response.split("[PART4]")[1].strip()
        return parts
    except Exception as e:
        return {"P1": f"오류: {e}", "P2": "", "P3": "", "P4": ""}


# 5. 메인 UI (파일 업로드)
uploaded_file = st.file_uploader("📂 분석할 PDF 파일을 업로드하세요", type=['pdf'])

if uploaded_file:
    if st.button("🔍 정밀 분석 시작"):
        with st.spinner("AI가 분석 중입니다..."):
            from pypdf import PdfReader
            reader = PdfReader(uploaded_file)
            full_text = "".join([p.extract_text() for p in reader.pages])
            st.session_state['result'] = run_analysis(uploaded_file.name,
                                                      full_text)

# 6. [그 상태 그대로] 결과 편집 및 4개 다운로드 UI
if 'result' in st.session_state:
    res = st.session_state['result']
    st.subheader("📝 파트별 편집 및 개별 저장")

    # 표 형태 레이아웃 (2x2)
    col1, col2 = st.columns(2)

    with col1:
        p1 = st.text_area("📊 Part 1. 성과 요약", value=res['P1'], height=250)
        st.download_button("📥 P1 다운로드", save_to_word("성과 요약", p1),
                           "Part1.docx")
        st.divider()
        p3 = st.text_area("💡 Part 3. 개선 계획", value=res['P3'], height=250)
        st.download_button("📥 P3 다운로드", save_to_word("개선 계획", p3),
                           "Part3.docx")

    with col2:
        p2 = st.text_area("⚠️ Part 2. 문제점", value=res['P2'], height=250)
        st.download_button("📥 P2 다운로드", save_to_word("문제점 및 한계", p2),
                           "Part2.docx")
        st.divider()
        p4 = st.text_area("🧐 Part 4. 총평", value=res['P4'], height=250)
        st.download_button("📥 P4 다운로드", save_to_word("총평", p4), "Part4.docx")

    st.divider()
    if st.button("📄 전체 통합 보고서 만들기"):
        total = f"1. 성과\n{p1}\n\n2. 문제점\n{p2}\n\n3. 계획\n{p3}\n\n4. 총평\n{p4}"
        st.download_button("📥 통합본 다운로드", save_to_word("종합 평가서", total),
                           "Total_Report.docx")
